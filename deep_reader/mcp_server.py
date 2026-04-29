"""MCP server exposing the deep-reader vault to Claude Desktop (or any MCP client).

Resources (read-only, loaded by the client):
  vault://summary                — top-level summary + recent activity
  vault://action_items           — central list (Nicole's to-dos)
  vault://waiting_on             — things owed to Nicole
  vault://people                 — index of people
  vault://people/{slug}          — a person page
  vault://sources/{slug}         — a source overview
  vault://threads/{name}         — a thread page
  vault://recaps/{date}          — a recap file
  vault://inbox                  — list of files in the inbox

Tools (actions):
  search(query, limit?)
  list_action_items(status?)
  list_waiting_on(person?, status?)
  add_action_item(description, source?)
  add_waiting_on(description, person, source?)
  close_action_item(id)
  list_people(query?)
  get_person(name)
  merge_people(keep, drop)
  list_inbox()
  ingest_file(filename, source_type?)
  ingest_file_bytes(content_base64, filename, mime_type, source_type?)
  ingest_note(text, title?)
  ingest_meeting(text, title?, date?, attendees?)
  recap_prep(date?)
  sync_recap(date?)

All tools return JSON-serializable dicts. Write tools accept MCP client
confirmation (Claude Desktop surfaces a confirmation UI by default).

Run with: `deep-reader mcp`
"""
from __future__ import annotations

import base64
import json
import sys
from datetime import date, datetime
from pathlib import Path
from typing_extensions import NotRequired, TypedDict

from deep_reader.config import Config, get_config
from deep_reader.llm import claude_code_llm
from deep_reader.search import search as search_fn
from deep_reader.state import GlobalState
from deep_reader.wiki import Wiki


# ---------- Nested structured types ----------
#
# Declared as TypedDicts so FastMCP's schema generator produces proper JSON
# Schema for each record_* tool. This gives Claude Desktop the exact shape it
# needs before calling (no guessing from docstrings) and turns validation
# errors into "expected key 'body' in thread_updates[0]" instead of the
# useless "KeyError: 'body'" that surfaced from downstream.

class Attendee(TypedDict):
    """Meeting attendee or doc author."""
    name: str
    role: NotRequired[str]
    email: NotRequired[str]


class ThreadUpdate(TypedDict):
    """Evidence to append to an existing thread.

    `slug` must match an existing thread (check via get_ingest_context).
    `body` is a ONE-sentence evidence entry, not a rewrite of the thesis.
    """
    slug: str
    body: str


class NewThread(TypedDict):
    """A new thread to create for a recurring theme worth tracking."""
    slug: str
    thesis: str


class PersonItem(TypedDict):
    """An item attributed to a named person (waiting_on / other_commitments)."""
    person: str
    description: str


def build_server(vault_root: Path):
    """Build an MCP server bound to the given vault.

    Uses the `mcp` Python SDK (FastMCP API).
    """
    try:
        from mcp.server.fastmcp import FastMCP
    except ImportError as e:
        raise SystemExit(
            "The `mcp` Python package is required. Install with: pip install mcp"
        ) from e

    config = get_config(vault_root)
    config.ensure_dirs()
    mcp = FastMCP("deep-reader")

    # ---------- Resources ----------

    @mcp.resource("vault://summary")
    def summary() -> str:
        state = GlobalState.load(config.state_file)
        wiki = Wiki(config)
        lines = [f"# Vault summary\n"]
        lines.append(f"- Sources: {len(state.sources)}")
        lines.append(f"- People: {len(state.people)}")
        lines.append(f"- Threads: {len(state.global_threads)}")
        open_mine = sum(
            1 for a in state.action_items if a.category == "mine" and a.status == "open"
        )
        open_waiting = sum(
            1 for a in state.action_items if a.category == "waiting_on" and a.status == "open"
        )
        pending_reviews = sum(1 for r in state.review_queue if r.status == "pending")
        lines.append(f"- Open action items: {open_mine}")
        lines.append(f"- Waiting on: {open_waiting}")
        if pending_reviews:
            lines.append(
                f"- **Pending reviews: {pending_reviews}** "
                f"(see `vault://review_pending` or run `/review_pending` in chat)"
            )
        if state.owner.name:
            lines.append(f"- Owner: {state.owner.name} <{state.owner.email}>")
        # Recent sources
        recent = sorted(
            [(s, src) for s, src in state.sources.items() if src.completed_at],
            key=lambda x: x[1].completed_at or datetime.min,
            reverse=True,
        )[:10]
        if recent:
            lines.append("\n## Recent sources")
            for slug, src in recent:
                when = src.completed_at.date().isoformat() if src.completed_at else "?"
                lines.append(f"- {slug} ({when})")
        g = wiki.read_summary()
        if g:
            lines.append("\n## Global summary\n" + g)
        return "\n".join(lines)

    @mcp.resource("vault://action_items")
    def action_items() -> str:
        if config.action_items_file.exists():
            return config.action_items_file.read_text()
        return "# My Action Items\n\n_(none yet)_\n"

    @mcp.resource("vault://waiting_on")
    def waiting_on() -> str:
        if config.waiting_on_file.exists():
            return config.waiting_on_file.read_text()
        return "# Waiting On\n\n_(none yet)_\n"

    @mcp.resource("vault://people")
    def people_index() -> str:
        path = config.wiki_indexes / "people.md"
        if path.exists():
            return path.read_text()
        return "# People\n\n_(none yet)_\n"

    @mcp.resource("vault://people/{slug}")
    def person_page(slug: str) -> str:
        path = config.wiki_people / f"{slug}.md"
        if not path.exists():
            return f"# {slug}\n\n_(no page)_\n"
        return path.read_text()

    @mcp.resource("vault://sources/{slug}")
    def source_page(slug: str) -> str:
        """Full source content: overview + every chunk page.

        For short sources (meetings, notes, short docs) this is the
        overview summary plus the single chunk-001.md, which contains
        decisions, attendees, action items, concepts — everything Claude
        parsed during record_*. For long-form sources (books, papers)
        it's the overview plus all chunks concatenated.

        This is what Claude should fetch after search() identifies a
        source as relevant — the search result only returns the first
        paragraph of the overview, which isn't enough for substantive
        questions.
        """
        src_dir = config.wiki_sources / slug
        if not src_dir.exists():
            return f"# {slug}\n\n_(source not found)_\n"
        parts: list[str] = []
        overview = src_dir / "_overview.md"
        if overview.exists():
            parts.append(overview.read_text())
        for chunk_path in sorted(src_dir.glob("chunk-*.md")):
            parts.append(f"\n---\n\n## {chunk_path.stem}\n")
            parts.append(chunk_path.read_text())
        return "\n".join(parts) if parts else f"# {slug}\n\n_(no content)_\n"

    @mcp.resource("vault://threads/{name}")
    def thread_page(name: str) -> str:
        path = config.wiki_threads / f"{name}.md"
        if not path.exists():
            return f"# {name}\n\n_(no thread)_\n"
        return path.read_text()

    @mcp.resource("vault://recaps/{when}")
    def recap_page(when: str) -> str:
        path = config.recaps / f"{when}.md"
        if not path.exists():
            return f"# Recap {when}\n\n_(not found)_\n"
        return path.read_text()

    @mcp.resource("vault://review_pending")
    def review_pending_resource() -> str:
        path = config.vault_root / "wiki" / "_review" / "pending.md"
        if path.exists():
            return path.read_text()
        return "# Pending Reviews\n\n_(nothing waiting)_\n"

    @mcp.resource("vault://inbox")
    def inbox_list() -> str:
        if not config.inbox.exists():
            return "# Inbox\n\n_(empty)_\n"
        files = sorted(config.inbox.iterdir())
        if not files:
            return "# Inbox\n\n_(empty)_\n"
        lines = ["# Inbox\n"]
        for f in files:
            if f.is_file():
                size = f.stat().st_size
                lines.append(f"- {f.name} ({size:,} bytes)")
        return "\n".join(lines)

    # ---------- Tools ----------

    # ---------- Index / routing tools ----------
    #
    # These compute relationships from existing state. They return
    # structural pointers (slugs, counts, dates) — never paraphrased
    # content. Claude uses them to figure out WHICH source/thread/person
    # to read, then loads the actual content via get_source / search.

    @mcp.tool()
    def find_related(slug: str, limit: int = 10) -> dict:
        """Return entities most-connected to the given slug.

        Auto-detects entity type (source / person / thread / concept). For:
          - source → co-attendees, threads it advanced, concepts tagged
          - person → sources they appear in, threads they're central to,
            people they co-appear with most
          - thread → contributing sources, central people, related threads
            (sharing 2+ sources), concepts the thread touches
          - concept → contributing sources, parent/child/related concepts,
            central people, threads advancing it
        """
        state = GlobalState.load(config.state_file)
        # Resolve type
        if slug in state.sources:
            return _related_for_source(state, slug, limit)
        if slug in state.people:
            return _related_for_person(state, slug, limit)
        if slug in state.global_threads:
            return _related_for_thread(state, slug, config, limit)
        if slug in state.concepts:
            return _related_for_concept(state, slug, config, limit)
        # Try concept by lowercase
        slug_l = slug.lower()
        if slug_l in state.concepts:
            return _related_for_concept(state, slug_l, config, limit)
        return {"error": f"slug '{slug}' not found as source/person/thread/concept"}

    @mcp.tool()
    def who_knows_about(topic: str, limit: int = 10) -> list[dict]:
        """Rank people by source-overlap with a topic (thread or concept).

        Resolves topic to a thread or concept, then counts each person's
        appearances in sources contributing to that topic.
        """
        state = GlobalState.load(config.state_file)
        # Resolve topic to set of contributing source slugs
        contributing: set[str] = set()
        topic_l = topic.lower().replace(" ", "-")
        if topic_l in state.global_threads:
            # Source slugs appearing in the thread's evidence
            thread_path = config.wiki_threads / f"{topic_l}.md"
            if thread_path.exists():
                from deep_reader.thread_utils import extract_section
                evidence = extract_section(thread_path.read_text(), "Evidence")
                import re as _re
                for m in _re.finditer(r"\[\[([^\]/]+)/chunk-\d+\]\]", evidence):
                    contributing.add(m.group(1))
        else:
            # Treat as concept — find sources tagging this concept
            import re as _re
            for src_slug in state.sources:
                src_dir = config.wiki_sources / src_slug
                for cp in src_dir.glob("chunk-*.md"):
                    page = cp.read_text()
                    if _re.search(r"\[\[" + _re.escape(topic_l) + r"\]\]", page, _re.IGNORECASE):
                        contributing.add(src_slug)
                        break
        if not contributing:
            return []
        # Rank people by overlap
        ranked = []
        for p in state.people.values():
            overlap = len(set(p.appearances) & contributing)
            if overlap:
                ranked.append({
                    "slug": p.slug, "name": p.name, "role": p.role,
                    "appearances_overlap": overlap,
                    "total_appearances": len(p.appearances),
                })
        ranked.sort(key=lambda x: x["appearances_overlap"], reverse=True)
        return ranked[:limit]

    @mcp.tool()
    def overlap(slug_a: str, slug_b: str) -> dict:
        """Shared sources / threads / concepts between two entities.

        Useful for 'what do Jane and Bob have in common' or 'what
        threads do these two people both appear in'.
        """
        state = GlobalState.load(config.state_file)
        sources_a = _entity_sources(state, slug_a, config)
        sources_b = _entity_sources(state, slug_b, config)
        if not sources_a or not sources_b:
            return {"error": "one or both slugs not found"}
        shared_sources = sorted(sources_a & sources_b)
        return {
            "a": slug_a,
            "b": slug_b,
            "shared_sources": shared_sources,
            "shared_count": len(shared_sources),
            "a_only_count": len(sources_a - sources_b),
            "b_only_count": len(sources_b - sources_a),
        }

    @mcp.tool()
    def timeline(
        person: str | None = None,
        thread: str | None = None,
        concept: str | None = None,
        since_days: int = 90,
        limit: int = 50,
    ) -> list[dict]:
        """Chronological event stream from the vault.

        Filters: by person (sources where they appear), by thread (sources
        in evidence), by concept (sources tagging it), by date window
        (since_days, default 90).
        """
        from datetime import timedelta
        state = GlobalState.load(config.state_file)
        cutoff = datetime.now() - timedelta(days=since_days)
        # Build candidate source set based on filters
        candidates: set[str] = set(state.sources.keys())
        if person:
            person_slug = person if person in state.people else None
            if not person_slug:
                for p in state.people.values():
                    if p.name.lower() == person.lower():
                        person_slug = p.slug
                        break
            if person_slug:
                candidates &= set(state.people[person_slug].appearances)
            else:
                return []
        if thread:
            thread_path = config.wiki_threads / f"{thread}.md"
            if thread_path.exists():
                from deep_reader.thread_utils import extract_section
                import re as _re
                evidence = extract_section(thread_path.read_text(), "Evidence")
                thread_sources = {m.group(1) for m in _re.finditer(r"\[\[([^\]/]+)/chunk-\d+\]\]", evidence)}
                candidates &= thread_sources
        if concept:
            concept_l = concept.lower()
            import re as _re
            concept_sources: set[str] = set()
            for src_slug in candidates:
                src_dir = config.wiki_sources / src_slug
                for cp in src_dir.glob("chunk-*.md"):
                    if _re.search(r"\[\[" + _re.escape(concept_l) + r"\]\]", cp.read_text(), _re.IGNORECASE):
                        concept_sources.add(src_slug)
                        break
            candidates &= concept_sources

        events = []
        for slug in candidates:
            src = state.sources[slug]
            if src.completed_at and src.completed_at >= cutoff:
                events.append({
                    "kind": "source_ingested",
                    "when": src.completed_at.isoformat(),
                    "slug": slug,
                    "type": src.source_type,
                })
        # Action items in window
        for a in state.action_items:
            if a.created_at >= cutoff:
                if person and a.owner != (state.people.get(person, None) and state.people[person].slug):
                    continue
                events.append({
                    "kind": "action_item_created",
                    "when": a.created_at.isoformat(),
                    "id": a.id, "owner": a.owner, "category": a.category,
                    "description": a.description[:80],
                })
            if a.completed_at and a.completed_at >= cutoff:
                events.append({
                    "kind": "action_item_closed",
                    "when": a.completed_at.isoformat(),
                    "id": a.id, "description": a.description[:80],
                })
        events.sort(key=lambda e: e["when"], reverse=True)
        return events[:limit]

    @mcp.tool()
    def coverage(slug: str) -> dict:
        """All sources, people, time range contributing to a thread or concept.

        Pure metadata — no content. Use to understand the SCOPE of an
        entity before deciding to read source content.
        """
        state = GlobalState.load(config.state_file)
        contributing_sources: list[str] = []
        contributing_people: list[str] = []
        # Resolve as thread first
        thread_path = config.wiki_threads / f"{slug}.md"
        if thread_path.exists():
            from deep_reader.thread_utils import extract_section
            import re as _re
            evidence = extract_section(thread_path.read_text(), "Evidence")
            seen: set[str] = set()
            for m in _re.finditer(r"\[\[([^\]/]+)/chunk-\d+\]\]", evidence):
                if m.group(1) not in seen:
                    seen.add(m.group(1))
                    contributing_sources.append(m.group(1))
        elif slug in state.concepts or slug.lower() in state.concepts:
            slug_l = slug.lower() if slug.lower() in state.concepts else slug
            import re as _re
            for src_slug in state.sources:
                src_dir = config.wiki_sources / src_slug
                for cp in src_dir.glob("chunk-*.md"):
                    if _re.search(r"\[\[" + _re.escape(slug_l) + r"\]\]", cp.read_text(), _re.IGNORECASE):
                        contributing_sources.append(src_slug)
                        break
            slug = slug_l
        else:
            return {"error": f"'{slug}' not found as thread or concept"}

        # People involved
        for p in state.people.values():
            if set(p.appearances) & set(contributing_sources):
                contributing_people.append(p.slug)

        # Time range
        dates = []
        for s in contributing_sources:
            src = state.sources.get(s)
            if src and src.completed_at:
                dates.append(src.completed_at)
        first = min(dates).date().isoformat() if dates else None
        last = max(dates).date().isoformat() if dates else None

        return {
            "slug": slug,
            "source_count": len(contributing_sources),
            "sources": contributing_sources,
            "person_count": len(contributing_people),
            "people": contributing_people,
            "first_appearance": first,
            "last_appearance": last,
        }

    @mcp.tool()
    def recent_activity(slug: str, since_days: int = 30) -> dict:
        """What's happened around an entity recently.

        Returns recent sources mentioning the entity, recent action items
        attributed to it, and any thread evidence added in the window.
        """
        return timeline(  # type: ignore[name-defined]
            person=slug if slug in GlobalState.load(config.state_file).people else None,
            thread=slug,
            concept=slug,
            since_days=since_days,
            limit=30,
        )

    @mcp.tool()
    def connections_between(slug_a: str, slug_b: str) -> dict:
        """Find the path / shared context linking two entities.

        E.g., 'how is Jane connected to the Duckbill thread' → returns
        the sources where Jane appears AND that contribute to duckbill,
        the people they share, the threads/concepts they both touch.
        """
        state = GlobalState.load(config.state_file)
        sources_a = _entity_sources(state, slug_a, config)
        sources_b = _entity_sources(state, slug_b, config)
        shared_sources = sorted(sources_a & sources_b)
        # People who appear in any shared source
        shared_people: set[str] = set()
        for s in shared_sources:
            for p in state.people.values():
                if s in p.appearances:
                    shared_people.add(p.slug)
        return {
            "a": slug_a,
            "b": slug_b,
            "shared_sources": shared_sources,
            "shared_people": sorted(shared_people),
            "connection_strength": len(shared_sources),
        }

    # ---------- Concept hierarchy tools ----------

    @mcp.tool()
    def link_concepts(parent: str, child: str, kind: str = "parent") -> dict:
        """Establish a relationship between two concepts.

        kind: 'parent' (child is-a parent) or 'related' (peer relationship).
        Both concepts will be created in state if they don't exist.
        """
        state = GlobalState.load(config.state_file)
        result = _do_link_concepts(config, state, parent=parent, child=child, kind=kind)
        return result

    @mcp.tool()
    def unlink_concepts(slug_a: str, slug_b: str) -> dict:
        """Remove all relationships between two concepts."""
        state = GlobalState.load(config.state_file)
        a_slug = slug_a.strip().lower().replace(" ", "-")
        b_slug = slug_b.strip().lower().replace(" ", "-")
        for slug in (a_slug, b_slug):
            if slug not in state.concepts:
                continue
            c = state.concepts[slug]
            other = b_slug if slug == a_slug else a_slug
            c.parent_concepts = [s for s in c.parent_concepts if s != other]
            c.child_concepts = [s for s in c.child_concepts if s != other]
            c.related_concepts = [s for s in c.related_concepts if s != other]
        state.save(config.state_file)
        return {"unlinked": [a_slug, b_slug]}

    @mcp.tool()
    def get_concept_with_hierarchy(name: str, depth: int = 2) -> dict:
        """Return a concept + its parent chain (recursive) + children + related.

        depth controls how many levels of parent/child to walk (default 2).
        Useful when you're discussing a child concept and want to bring in
        its parents' contextual content.
        """
        state = GlobalState.load(config.state_file)
        slug = name.strip().lower().replace(" ", "-")
        if slug not in state.concepts:
            return {"error": f"concept '{name}' not found"}

        # Walk parent chain up to `depth`
        parent_chain: list[str] = []
        seen: set[str] = set()
        frontier = [slug]
        for _level in range(depth):
            next_frontier = []
            for s in frontier:
                if s not in state.concepts:
                    continue
                for p in state.concepts[s].parent_concepts:
                    if p not in seen and p != slug:
                        seen.add(p)
                        parent_chain.append(p)
                        next_frontier.append(p)
            frontier = next_frontier
            if not frontier:
                break

        c = state.concepts[slug]
        return {
            "slug": slug,
            "name": c.name,
            "parents": parent_chain,
            "children": list(c.child_concepts),
            "related": list(c.related_concepts),
            "sources_at_last_refresh": c.sources_at_last_refresh,
            "last_refreshed": c.last_refreshed.isoformat() if c.last_refreshed else None,
            "page_exists": (config.wiki_concepts / f"{slug}.md").exists(),
        }

    @mcp.tool()
    def list_stale_concepts(min_new_sources: int = 3) -> list[dict]:
        """Concepts whose page is out of date relative to current source coverage.

        Returns concepts where (current source-tag count) - (sources at last
        refresh) >= threshold, OR concepts with no page yet that have >= 3
        source tags.
        """
        state = GlobalState.load(config.state_file)
        # Compute current source-tag count per concept
        from collections import defaultdict
        import re as _re
        coverage_now: dict[str, set[str]] = defaultdict(set)
        for src_slug in state.sources:
            src_dir = config.wiki_sources / src_slug
            for cp in src_dir.glob("chunk-*.md"):
                page = cp.read_text()
                in_concepts = False
                buf = []
                for line in page.split("\n"):
                    if line.startswith("## Concepts"):
                        in_concepts = True
                        continue
                    if in_concepts and line.startswith("## "):
                        break
                    if in_concepts:
                        buf.append(line)
                section = "\n".join(buf)
                for m in _re.finditer(r"\[\[([^\]]+)\]\]", section):
                    coverage_now[m.group(1).strip().lower()].add(src_slug)

        stale: list[dict] = []
        for name, sources in coverage_now.items():
            current_count = len(sources)
            concept = state.concepts.get(name)
            page_exists = (config.wiki_concepts / f"{name}.md").exists()

            if not page_exists and current_count >= 3:
                stale.append({
                    "name": name,
                    "reason": "no_page_yet",
                    "current_sources": current_count,
                    "sources_at_last_refresh": 0,
                })
                continue
            if concept and (current_count - concept.sources_at_last_refresh) >= min_new_sources:
                stale.append({
                    "name": name,
                    "reason": "new_sources_since_refresh",
                    "current_sources": current_count,
                    "sources_at_last_refresh": concept.sources_at_last_refresh,
                    "delta": current_count - concept.sources_at_last_refresh,
                })
        stale.sort(key=lambda x: x.get("delta", x["current_sources"]), reverse=True)
        return stale

    @mcp.tool()
    def record_concept_page(
        name: str,
        definition: str,
        distillation: str,
        contributing_sources: list[str],
        parent_concepts: list[str] | None = None,
        child_concepts: list[str] | None = None,
        related_concepts: list[str] | None = None,
        tensions: str = "",
    ) -> dict:
        """Write or replace a concept page at /wiki/concepts/{slug}.md.

        Concept pages are the ONE place persistent prose synthesis is
        appropriate — concepts are meta-entities that exist as integrations
        across sources. Constraints:
          - definition: 1-paragraph definition in this vault's working context
          - distillation: 2-3 paragraph integrated reflection. Heavy citations
            via [[<source-slug>]]. Direct quotes where they sharpen the
            picture. Not a paraphrase that replaces source material.
          - contributing_sources: every source slug tagging this concept
          - hierarchy fields: optional, link concepts via link_concepts too

        For initial creation OR refresh; both go through this tool. The
        existing page (if any) is replaced — but typically called via the
        review queue (`propose_review` with kind='concept_refresh') so the
        user sees a diff first.
        """
        state = GlobalState.load(config.state_file)
        return _do_record_concept_page(
            config, state,
            name=name, definition=definition, distillation=distillation,
            contributing_sources=contributing_sources,
            parent_concepts=parent_concepts or [],
            child_concepts=child_concepts or [],
            related_concepts=related_concepts or [],
            tensions=tensions,
        )

    # ---------- Drive tracking ----------

    @mcp.tool()
    def is_drive_ingested(drive_id: str) -> dict:
        """Check if a Drive doc has already been ingested.

        Returns the source slug if yes, null if no. Use during /crawl_drive
        to skip docs already in the vault.
        """
        state = GlobalState.load(config.state_file)
        return {
            "drive_id": drive_id,
            "ingested": drive_id in state.drive.ingested_ids,
            "source_slug": state.drive.ingested_ids.get(drive_id),
        }

    @mcp.tool()
    def mark_drive_ingested(drive_id: str, source_slug: str) -> dict:
        """Record that a Drive doc has been ingested as a given source.

        Call this after record_doc completes for a Drive-sourced document.
        """
        state = GlobalState.load(config.state_file)
        state.drive.ingested_ids[drive_id] = source_slug
        state.drive.last_crawl_at = datetime.now()
        state.save(config.state_file)
        return {"drive_id": drive_id, "source_slug": source_slug, "tracked": True}

    @mcp.tool()
    def list_drive_ingested() -> dict:
        """All Drive doc IDs currently tracked as ingested."""
        state = GlobalState.load(config.state_file)
        return {
            "count": len(state.drive.ingested_ids),
            "last_crawl_at": state.drive.last_crawl_at.isoformat() if state.drive.last_crawl_at else None,
            "ids": dict(state.drive.ingested_ids),
        }

    # ---------- Review queue ----------
    #
    # Some Claude-proposed actions need user approval before being made
    # final — concept page replacements, hierarchy suggestions, Drive
    # ingest candidates, borderline-relevance docs. Those go into a
    # persistent review queue rendered to /wiki/_review/pending.md and
    # surfaced via the vault summary.

    @mcp.tool()
    def list_pending_reviews(kind: str | None = None) -> list[dict]:
        """List items waiting for user approval. Optionally filter by kind."""
        state = GlobalState.load(config.state_file)
        items = [r for r in state.review_queue if r.status == "pending"]
        if kind:
            items = [r for r in items if r.kind == kind]
        items.sort(key=lambda r: r.created_at, reverse=True)
        return [_dump_review(r) for r in items]

    @mcp.tool()
    def get_review(id: str) -> dict:
        """Get full details of a single review item, including its proposed action."""
        state = GlobalState.load(config.state_file)
        for r in state.review_queue:
            if r.id == id:
                return _dump_review(r, full=True)
        return {"error": f"no review with id '{id}'"}

    @mcp.tool()
    def propose_review(
        kind: str,
        title: str,
        preview: str,
        proposed_action: dict,
    ) -> dict:
        """Queue an action for user approval.

        kind: one of `concept_refresh`, `concept_link`, `enrichment_ingest`,
              `drive_borderline`, or any other label Claude finds useful.
        title: a short one-line summary the user will scan.
        preview: a longer multi-line description / diff so the user can
              decide informed.
        proposed_action: {tool: str, args: dict} — what gets executed on
              approve. Must reference an existing tool name.

        Returns the assigned ID so Claude can mention it back to the user.
        """
        import hashlib as _hash
        state = GlobalState.load(config.state_file)
        rid = _hash.sha1(
            f"{kind}|{title}|{datetime.now().isoformat()}".encode()
        ).hexdigest()[:10]
        item = state.review_queue.__class__.__args__[0] if False else None
        # Direct construct since the type isn't easily importable from here
        from deep_reader.state import ReviewItem
        item = ReviewItem(
            id=rid,
            kind=kind,
            title=title,
            preview=preview,
            proposed_action=proposed_action,
            created_at=datetime.now(),
        )
        state.review_queue.append(item)
        state.save(config.state_file)
        _render_review_pending(config, state)
        return {"id": rid, "queued": True}

    @mcp.tool()
    def approve_review(id: str) -> dict:
        """Approve a queued action — executes the proposed_action and marks done.

        Dispatches to the named tool with the stored args. Returns the
        result of the executed tool, or an error if the tool isn't
        recognized / doesn't exist.
        """
        state = GlobalState.load(config.state_file)
        item = next((r for r in state.review_queue if r.id == id), None)
        if not item:
            return {"error": f"no review with id '{id}'"}
        if item.status != "pending":
            return {"error": f"review {id} already {item.status}"}

        # Execute the proposed action by dispatching to the named tool.
        action = item.proposed_action or {}
        tool_name = action.get("tool")
        tool_args = action.get("args", {})
        result = _dispatch_review_action(config, tool_name, tool_args)

        # Mark the item resolved
        item.status = "approved"
        item.reviewed_at = datetime.now()
        state.save(config.state_file)
        _render_review_pending(config, state)
        return {"id": id, "status": "approved", "execution_result": result}

    @mcp.tool()
    def reject_review(id: str, reason: str = "") -> dict:
        """Reject a queued action without executing it."""
        state = GlobalState.load(config.state_file)
        item = next((r for r in state.review_queue if r.id == id), None)
        if not item:
            return {"error": f"no review with id '{id}'"}
        if item.status != "pending":
            return {"error": f"review {id} already {item.status}"}
        item.status = "rejected"
        item.reviewed_at = datetime.now()
        state.save(config.state_file)
        _render_review_pending(config, state)
        return {"id": id, "status": "rejected", "reason": reason}

    @mcp.tool()
    def search(
        query: str,
        limit: int = 10,
        depth: str = "full",
        inline_top_n: int = 5,
    ) -> dict:
        """Search the vault. Returns FULL content of top hits by default.

        depth="full" (default): inline the full content of the top
        `inline_top_n` source hits and top `inline_top_n` thread hits,
        so Claude can answer substantive questions from a single call
        without follow-up retrieval. Default 5 — bump higher (7-10) when
        you have a wider vault and want broader synthesis, lower (3) for
        tight token budgets.

        depth="lite": routing snippets only.

        Fields returned:
          sources: [{slug, score, snippet, content?}]
          threads: [{name, score, thesis, content?}]
          concepts, people, action_items: structured metadata
        """
        r = search_fn(query, config, limit=limit)
        out: dict = {
            "sources": list(r.sources),
            "threads": list(r.threads),
            "concepts": r.concepts,
            "people": r.people,
            "action_items": r.action_items,
            "depth": depth,
            "inline_top_n": inline_top_n if depth == "full" else 0,
        }

        if depth == "full":
            n = max(1, min(inline_top_n, 20))  # bound it
            for i, src in enumerate(out["sources"][:n]):
                slug = src["slug"]
                src_dir = config.wiki_sources / slug
                parts = []
                ov = src_dir / "_overview.md"
                if ov.exists():
                    parts.append(ov.read_text())
                for cp in sorted(src_dir.glob("chunk-*.md")):
                    parts.append(f"\n---\n\n## {cp.stem}\n" + cp.read_text())
                out["sources"][i] = {**src, "content": "\n".join(parts)}
            for i, t in enumerate(out["threads"][:n]):
                path = config.wiki_threads / f"{t['name']}.md"
                if path.exists():
                    out["threads"][i] = {**t, "content": path.read_text()}

        return out

    @mcp.tool()
    def get_source(slug: str) -> dict:
        """Return the full content of a source: overview + all chunk pages.

        Use this after `search` identifies a source as relevant. The
        search tool only returns the first paragraph of the overview —
        this tool gives you everything, including the structured sections
        (decisions, attendees, action items, concepts) that were parsed
        during ingest.
        """
        src_dir = config.wiki_sources / slug
        if not src_dir.exists():
            return {"error": f"source '{slug}' not found"}
        overview_path = src_dir / "_overview.md"
        overview = overview_path.read_text() if overview_path.exists() else ""
        chunks = []
        for chunk_path in sorted(src_dir.glob("chunk-*.md")):
            chunks.append({
                "name": chunk_path.stem,
                "content": chunk_path.read_text(),
            })
        state = GlobalState.load(config.state_file)
        src_state = state.sources.get(slug)
        meta: dict = {"slug": slug}
        if src_state:
            meta["source_type"] = src_state.source_type
            meta["attendees"] = src_state.attendees
            if src_state.meeting_date:
                meta["date"] = src_state.meeting_date
        return {
            "meta": meta,
            "overview": overview,
            "chunks": chunks,
        }

    @mcp.tool()
    def list_action_items(status: str = "open") -> list[dict]:
        """List the vault owner's action items (category=mine)."""
        state = GlobalState.load(config.state_file)
        items = [a for a in state.action_items if a.category == "mine"]
        if status != "all":
            items = [a for a in items if a.status == status]
        items.sort(key=lambda a: a.created_at)
        return [_dump_action(a) for a in items]

    @mcp.tool()
    def list_waiting_on(person: str | None = None, status: str = "open") -> list[dict]:
        """List waiting-on items, optionally filtered by person slug or name."""
        state = GlobalState.load(config.state_file)
        items = [a for a in state.action_items if a.category == "waiting_on"]
        if status != "all":
            items = [a for a in items if a.status == status]
        if person:
            from deep_reader.steps.people import slugify_name
            target = slugify_name(person)
            items = [a for a in items if a.owner == target or a.owner == person]
        items.sort(key=lambda a: (a.owner, a.created_at))
        return [_dump_action(a, state) for a in items]

    @mcp.tool()
    def add_action_item(description: str, source: str = "chat") -> dict:
        """Add a new personal action item (owned by the vault owner)."""
        from deep_reader.steps import actions as actions_step
        state = GlobalState.load(config.state_file)
        item = actions_step.add_mine(state, description, source)
        state.save(config.state_file)
        _rerender_after_action_change(config, state, item)
        return _dump_action(item)

    @mcp.tool()
    def add_waiting_on(description: str, person: str, source: str = "chat") -> dict:
        """Add a waiting-on item owed by a specific person."""
        from deep_reader.steps import actions as actions_step
        state = GlobalState.load(config.state_file)
        item = actions_step.add_waiting_on(state, description, person, source)
        state.save(config.state_file)
        _rerender_after_action_change(config, state, item)
        return _dump_action(item, state)

    @mcp.tool()
    def forget_source(source_slug: str) -> dict:
        """Remove a source from the vault: its page, state, attributed action
        items, and any thread evidence referencing it. The raw file in raw/
        is preserved. People records are NOT deleted (they may appear in
        other sources) — this just decrements their appearance list.
        """
        import shutil
        from deep_reader.thread_utils import extract_section, assemble_thread
        state = GlobalState.load(config.state_file)
        wiki = Wiki(config)
        if source_slug not in state.sources:
            return {"error": f"no source with slug '{source_slug}'"}

        src_dir = wiki.source_dir(source_slug)
        if src_dir.exists():
            shutil.rmtree(src_dir)

        removed_actions = [a for a in state.action_items if a.source == source_slug]
        state.action_items = [a for a in state.action_items if a.source != source_slug]

        evidence_ref = f"[[{source_slug}/"
        threads_touched = []
        for thread_path in config.wiki_threads.glob("*.md"):
            content = thread_path.read_text()
            evidence = extract_section(content, "Evidence")
            if evidence_ref not in evidence:
                continue
            new_lines = [line for line in evidence.split("\n") if evidence_ref not in line]
            new_evidence = "\n".join(new_lines).strip() or "(no evidence yet)"
            thesis = extract_section(content, "Thesis")
            status = extract_section(content, "Status")
            thread_path.write_text(assemble_thread(thesis, new_evidence, status))
            threads_touched.append(thread_path.stem)

        affected_owners = {a.owner for a in removed_actions}
        people_with_removed_appearance = [
            p for p in state.people.values() if source_slug in p.appearances
        ]
        for p in people_with_removed_appearance:
            p.appearances.remove(source_slug)

        del state.sources[source_slug]
        state.save(config.state_file)

        from deep_reader.wiki import render_action_items, render_waiting_on
        from deep_reader.steps import people as people_step
        render_action_items(wiki, state)
        render_waiting_on(wiki, state)

        # Re-render every person page whose appearance list shrank or whose
        # waiting-on items were part of this source.
        to_refresh = {p.slug for p in people_with_removed_appearance} | affected_owners
        for p in state.people.values():
            if state.owner.matches(p.name) or state.owner.matches(p.email or ""):
                to_refresh.add(p.slug)
        for slug in to_refresh:
            if slug in state.people:
                people_step.render_person_page(
                    state.people[slug], state, config.wiki_people
                )

        return {
            "forgotten": source_slug,
            "action_items_removed": len(removed_actions),
            "threads_scrubbed": threads_touched,
            "people_pages_refreshed": len(to_refresh),
        }

    @mcp.tool()
    def close_action_item(id: str) -> dict:
        """Mark an action item (mine or waiting-on) as done."""
        from deep_reader.steps import actions as actions_step
        state = GlobalState.load(config.state_file)
        item = actions_step.close(state, id)
        if not item:
            return {"error": f"no item with id {id}"}
        state.save(config.state_file)
        _rerender_after_action_change(config, state, item)
        return _dump_action(item, state)

    # ---------- Synthesis: context-fetch + write-back tools ----------
    #
    # As the vault grows, the right shape for chat shifts from "Claude reads
    # raw chunks on demand" to "Claude reads continuously-maintained
    # synthesis articles". These tools support that shift: get_*_context
    # bundles the source material Claude needs to write a synthesis;
    # update_*_thesis / update_*_summary persist the result back. Pair with
    # the /refresh_* prompts.

    @mcp.tool()
    def get_thread_full_context(slug: str, max_evidence: int = 30) -> dict:
        """Return a thread + the content of every source it has evidence in.

        Use before regenerating a thread's thesis. Bundles:
          - the current thread file (thesis, evidence, status)
          - for each source referenced in the evidence section: the
            source's overview + chunk-001 (or all chunks if multi-chunk),
            up to max_evidence sources.

        This is what /refresh_thread_synthesis calls to get everything
        needed to write a richer thesis in one tool call.
        """
        from deep_reader.thread_utils import extract_section
        path = config.wiki_threads / f"{slug}.md"
        if not path.exists():
            return {"error": f"thread '{slug}' not found"}
        thread_content = path.read_text()
        thesis = extract_section(thread_content, "Thesis")
        evidence = extract_section(thread_content, "Evidence")
        status = extract_section(thread_content, "Status")

        # Parse evidence lines for source references like [[<slug>/chunk-NNN]]
        import re as _re
        source_slugs: list[str] = []
        seen: set[str] = set()
        for match in _re.finditer(r"\[\[([^\]/]+)/chunk-\d+\]\]", evidence):
            slug_ref = match.group(1)
            if slug_ref not in seen:
                seen.add(slug_ref)
                source_slugs.append(slug_ref)

        sources: list[dict] = []
        for src_slug in source_slugs[:max_evidence]:
            src_dir = config.wiki_sources / src_slug
            if not src_dir.exists():
                continue
            parts = []
            ov = src_dir / "_overview.md"
            if ov.exists():
                parts.append(ov.read_text())
            for cp in sorted(src_dir.glob("chunk-*.md")):
                parts.append(f"\n---\n\n## {cp.stem}\n" + cp.read_text())
            sources.append({"slug": src_slug, "content": "\n".join(parts)})

        return {
            "slug": slug,
            "current_thesis": thesis,
            "evidence_log": evidence,
            "status": status,
            "evidence_source_count": len(source_slugs),
            "evidence_sources_returned": len(sources),
            "sources": sources,
        }

    @mcp.tool()
    def get_person_full_context(slug: str, max_sources: int = 30) -> dict:
        """Return a person + the content of every source they appear in.

        Used before regenerating a person's summary. Returns:
          - the Person record (name, role, email, aliases, appearance count)
          - for each source they appear in: the source's full content,
            up to max_sources.

        Also returns recent action items they own (waiting-on items the
        vault owner is owed) for context.
        """
        state = GlobalState.load(config.state_file)
        if slug not in state.people:
            # Try resolving by name
            for p in state.people.values():
                if p.name.lower() == slug.lower():
                    slug = p.slug
                    break
            else:
                return {"error": f"person '{slug}' not found"}
        person = state.people[slug]

        sources: list[dict] = []
        for src_slug in person.appearances[-max_sources:]:
            src_dir = config.wiki_sources / src_slug
            if not src_dir.exists():
                continue
            parts = []
            ov = src_dir / "_overview.md"
            if ov.exists():
                parts.append(ov.read_text())
            for cp in sorted(src_dir.glob("chunk-*.md")):
                parts.append(f"\n---\n\n## {cp.stem}\n" + cp.read_text())
            sources.append({"slug": src_slug, "content": "\n".join(parts)})

        open_waiting = [
            _dump_action(a, state)
            for a in state.action_items
            if a.owner == slug and a.status == "open" and a.category == "waiting_on"
        ]

        return {
            "slug": person.slug,
            "name": person.name,
            "email": person.email,
            "role": person.role,
            "aliases": person.aliases,
            "current_summary": person.summary,
            "total_appearances": len(person.appearances),
            "new_appearances_since_summary": person.new_appearances_since_summary,
            "first_seen": person.first_seen.isoformat() if person.first_seen else None,
            "last_seen": person.last_seen.isoformat() if person.last_seen else None,
            "sources": sources,
            "open_waiting_on_them": open_waiting,
        }

    @mcp.tool()
    def list_concept_candidates(min_sources: int = 3) -> list[dict]:
        """List concepts that appear in min_sources+ sources but don't yet
        have a graduated concept article (or whose article is stale).

        These are candidates for /compile_concepts to synthesize into
        articles in /wiki/concepts/.
        """
        # Scan all chunk pages for [[concept-name]] in ## Concepts sections
        state = GlobalState.load(config.state_file)
        wiki = Wiki(config)
        from collections import defaultdict
        import re as _re
        coverage: dict[str, set[str]] = defaultdict(set)
        for slug in state.sources:
            src_dir = config.wiki_sources / slug
            for cp in src_dir.glob("chunk-*.md"):
                page = cp.read_text()
                # Find ## Concepts section
                in_concepts = False
                buf = []
                for line in page.split("\n"):
                    if line.startswith("## Concepts"):
                        in_concepts = True
                        continue
                    if in_concepts and line.startswith("## "):
                        break
                    if in_concepts:
                        buf.append(line)
                section = "\n".join(buf)
                for m in _re.finditer(r"\[\[([^\]]+)\]\]", section):
                    coverage[m.group(1).strip().lower()].add(slug)

        candidates = []
        for name, sources in coverage.items():
            if len(sources) >= min_sources:
                article_path = config.wiki_concepts / f"{name}.md"
                candidates.append({
                    "name": name,
                    "source_count": len(sources),
                    "sources": sorted(sources),
                    "has_article": article_path.exists(),
                })
        candidates.sort(key=lambda c: c["source_count"], reverse=True)
        return candidates

    @mcp.tool()
    def get_concept_evidence(name: str) -> dict:
        """Return all sources that tag this concept + their content.

        Used before writing or refreshing a concept article. Returns
        every source page where `[[<name>]]` appears in its Concepts
        section, with full content.
        """
        state = GlobalState.load(config.state_file)
        import re as _re
        target = name.strip().lower()
        matching_sources: list[dict] = []
        for slug in state.sources:
            src_dir = config.wiki_sources / slug
            for cp in src_dir.glob("chunk-*.md"):
                page = cp.read_text()
                in_concepts = False
                buf = []
                for line in page.split("\n"):
                    if line.startswith("## Concepts"):
                        in_concepts = True
                        continue
                    if in_concepts and line.startswith("## "):
                        break
                    if in_concepts:
                        buf.append(line)
                section = "\n".join(buf)
                if _re.search(r"\[\[" + _re.escape(target) + r"\]\]", section, _re.IGNORECASE):
                    parts = []
                    ov = src_dir / "_overview.md"
                    if ov.exists():
                        parts.append(ov.read_text())
                    parts.append(f"\n---\n\n## {cp.stem}\n" + page)
                    matching_sources.append({
                        "slug": slug,
                        "content": "\n".join(parts),
                    })
                    break
        article_path = config.wiki_concepts / f"{name}.md"
        existing_article = article_path.read_text() if article_path.exists() else ""
        return {
            "name": name,
            "source_count": len(matching_sources),
            "sources": matching_sources,
            "existing_article": existing_article,
        }

    @mcp.tool()
    def link_action_item(id: str, source_ref: str) -> dict:
        """Attach an additional source reference to an existing action item.

        Use this when you've spotted a paraphrase / re-mention — e.g.,
        a Slack message reaffirming a commitment that was already captured
        from a meeting. Adds the new source (Slack permalink, source slug,
        URL, etc.) to the item's additional_sources without creating a
        duplicate. Preserves the full trail without inflating the list.
        """
        from deep_reader.steps import actions as actions_step
        state = GlobalState.load(config.state_file)
        item = actions_step.link_source(state, id, source_ref)
        if not item:
            return {"error": f"no item with id {id}"}
        state.save(config.state_file)
        _rerender_after_action_change(config, state, item)
        return _dump_action(item, state)

    @mcp.tool()
    def list_people(query: str | None = None) -> list[dict]:
        """List known people, optionally filtered."""
        state = GlobalState.load(config.state_file)
        rows = list(state.people.values())
        if query:
            q = query.lower()
            rows = [p for p in rows if q in p.name.lower() or q in (p.email or "").lower()]
        rows.sort(key=lambda p: p.name.lower())
        return [
            {
                "slug": p.slug, "name": p.name, "email": p.email,
                "role": p.role, "aliases": p.aliases,
                "appearances": len(p.appearances),
            }
            for p in rows
        ]

    @mcp.tool()
    def get_person(name: str) -> dict:
        """Return the full person record and rendered page."""
        from deep_reader.steps.people import slugify_name
        state = GlobalState.load(config.state_file)
        slug = slugify_name(name)
        person = state.people.get(slug)
        if not person:
            for p in state.people.values():
                if p.name.lower() == name.lower() or name.lower() in [a.lower() for a in p.aliases]:
                    person = p
                    break
        if not person:
            return {"error": f"no person matching '{name}'"}
        page_path = config.wiki_people / f"{person.slug}.md"
        page = page_path.read_text() if page_path.exists() else ""
        return {
            "slug": person.slug, "name": person.name, "email": person.email,
            "role": person.role, "aliases": person.aliases,
            "appearances": person.appearances,
            "page": page,
        }

    @mcp.tool()
    def merge_people(keep: str, drop: str) -> dict:
        """Merge `drop` person into `keep` (slugs)."""
        from deep_reader.steps.people import merge_people as merge_fn, render_all_people, render_people_index
        state = GlobalState.load(config.state_file)
        merged = merge_fn(state, keep, drop)
        state.save(config.state_file)
        render_all_people(state, config.wiki_people)
        render_people_index(state, config.wiki_indexes / "people.md")
        return {"slug": merged.slug, "name": merged.name, "aliases": merged.aliases}

    @mcp.tool()
    def list_inbox() -> list[dict]:
        """List files sitting in the vault inbox waiting to be ingested."""
        if not config.inbox.exists():
            return []
        out = []
        for f in sorted(config.inbox.iterdir()):
            if f.is_file():
                out.append({
                    "name": f.name,
                    "size": f.stat().st_size,
                    "suffix": f.suffix.lower(),
                })
        return out

    @mcp.tool()
    def ingest_file(filename: str, source_type: str | None = None) -> dict:
        """LEGACY: ingest a file from inbox with server-side LLM analysis.

        Requires ANTHROPIC_API_KEY. For the default workflow, use
        read_inbox_file + record_meeting/record_note/record_doc +
        move_inbox_file instead — no API key needed.
        """
        import os
        if not os.environ.get("ANTHROPIC_API_KEY"):
            return {
                "error": (
                    "ingest_file requires ANTHROPIC_API_KEY. Use the "
                    "no-API-key flow: read_inbox_file(filename) → "
                    "record_meeting/record_note/record_doc → "
                    "move_inbox_file(filename, type)."
                )
            }
        src = config.inbox / filename
        if not src.exists():
            return {"error": f"inbox/{filename} not found"}
        stype = source_type or _auto_detect_type_path(src)
        _ingest_path(config, src, stype)
        dest_dir = _raw_dir_for(config, stype)
        dest_dir.mkdir(parents=True, exist_ok=True)
        final = dest_dir / src.name
        src.rename(final)
        _read_new_source(config, final, stype)
        return {"ingested": str(final), "source_type": stype}

    @mcp.tool()
    def ingest_file_bytes(
        content_base64: str,
        filename: str,
        mime_type: str | None = None,
        source_type: str | None = None,
    ) -> dict:
        """LEGACY: ingest an inline base64 file with server-side LLM.

        Requires ANTHROPIC_API_KEY. Prefer the no-API-key flow.
        """
        import os
        if not os.environ.get("ANTHROPIC_API_KEY"):
            return {"error": "ingest_file_bytes requires ANTHROPIC_API_KEY."}
        data = base64.b64decode(content_base64)
        tmp = config.inbox / filename
        config.inbox.mkdir(parents=True, exist_ok=True)
        tmp.write_bytes(data)
        return ingest_file(filename, source_type)

    # ---------- NEW: structured ingest (no LLM on server side) ----------
    #
    # These tools accept data Claude Desktop has already analyzed. Our server
    # just persists it — no API key, no server-side LLM calls. This is the
    # primary flow for users running everything through their own Claude
    # account. See the MCP prompts below for the exact schema Claude must
    # produce to call these.

    @mcp.tool()
    def get_ingest_context() -> dict:
        """Return everything Claude needs to parse a new source well.

        Includes: vault owner identity (for the mine/waiting-on split), all
        active threads with their theses (so new sources can extend them),
        and all known people with aliases (so attendees resolve correctly).
        Call this first before analyzing a new source.
        """
        state = GlobalState.load(config.state_file)
        wiki = Wiki(config)
        threads = []
        for slug in state.global_threads:
            content = wiki.read_thread(slug) or ""
            from deep_reader.thread_utils import extract_section
            thesis = extract_section(content, "Thesis") if content else ""
            threads.append({"slug": slug, "thesis": thesis})
        people = [
            {
                "slug": p.slug, "name": p.name, "email": p.email,
                "role": p.role, "aliases": p.aliases,
            }
            for p in state.people.values()
        ]
        return {
            "owner": {
                "name": state.owner.name,
                "email": state.owner.email,
                "aliases": state.owner.aliases,
            },
            "threads": threads,
            "people": people,
            "source_count": len(state.sources),
        }

    @mcp.tool()
    def read_inbox_file(filename: str) -> dict:
        """Return the text content of a file sitting in vault/inbox/.

        Handles PDF, .docx, .md, .txt, .rtf. Returns the extracted text for
        Claude to analyze — does NOT move the file out of the inbox. Use
        record_meeting / record_note / record_doc to persist, then call
        move_inbox_file to archive the original.
        """
        src = config.inbox / filename
        if not src.exists():
            return {"error": f"inbox/{filename} not found"}
        suffix = src.suffix.lower()
        try:
            if suffix in {".md", ".txt", ".rtf"}:
                text = src.read_text(errors="replace")
            elif suffix == ".pdf":
                from deep_reader.sources.pdf import extract_pdf
                text = extract_pdf(src)
            elif suffix == ".docx":
                try:
                    import docx
                except ImportError:
                    return {"error": "python-docx not installed (pip install 'deep-reader[docx]')"}
                doc = docx.Document(str(src))
                text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
            else:
                return {"error": f"unsupported file type: {suffix}"}
        except Exception as e:
            return {"error": f"failed to read {filename}: {e}"}
        return {
            "filename": filename,
            "suffix": suffix,
            "size": src.stat().st_size,
            "text": text,
            "suggested_type": _auto_detect_type_path(src),
        }

    @mcp.tool()
    def move_inbox_file(filename: str, source_type: str) -> dict:
        """Archive a processed inbox file into raw/{type}/ after ingest.

        Call this after record_meeting/record_note/record_doc succeeds, so
        the inbox stays clean and the original file is preserved alongside
        the compiled wiki page.
        """
        src = config.inbox / filename
        if not src.exists():
            return {"error": f"inbox/{filename} not found"}
        dest_dir = _raw_dir_for(config, source_type)
        dest_dir.mkdir(parents=True, exist_ok=True)
        dest = dest_dir / filename
        src.rename(dest)
        return {"archived": str(dest.relative_to(config.vault_root))}

    @mcp.tool()
    def record_meeting(
        title: str,
        body: str,
        summary: str,
        attendees: list[Attendee] | None = None,
        decisions: list[str] | None = None,
        action_items_mine: list[str] | None = None,
        waiting_on: list[PersonItem] | None = None,
        other_commitments: list[PersonItem] | None = None,
        thread_updates: list[ThreadUpdate] | None = None,
        new_threads: list[NewThread] | None = None,
        concepts: list[str] | None = None,
        date: str | None = None,
    ) -> dict:
        """Persist a meeting Claude has already analyzed.

        Nested types (see schema in the tool's inputSchema):
          - attendees: [{name, role?, email?}]
          - decisions: [str]
          - action_items_mine: [str] — items owned by the vault owner
          - waiting_on: [{person, description}] — owed TO the owner
          - other_commitments: [{person, description}] — between other parties
          - thread_updates: [{slug, body}] — existing-thread evidence
          - new_threads: [{slug, thesis}] — brand-new threads
          - concepts: [str] — concept slugs
          - date: YYYY-MM-DD
        """
        return _do_structured_record(
            config=config, source_type="meeting",
            title=title, body=body, date=date,
            attendees=attendees or [],
            result_payload={
                "summary": summary,
                "attendees": attendees or [],
                "decisions": decisions or [],
                "action_items_mine": action_items_mine or [],
                "waiting_on": waiting_on or [],
                "other_commitments": other_commitments or [],
                "thread_updates": thread_updates or [],
                "new_threads": new_threads or [],
                "concepts": concepts or [],
            },
        )

    @mcp.tool()
    def record_note(
        title: str,
        body: str,
        summary: str,
        action_items_mine: list[str] | None = None,
        waiting_on: list[PersonItem] | None = None,
        thread_updates: list[ThreadUpdate] | None = None,
        new_threads: list[NewThread] | None = None,
        concepts: list[str] | None = None,
    ) -> dict:
        """Persist a short note Claude has already analyzed.

        Nested types (see schema in the tool's inputSchema):
          - action_items_mine: [str] — items owned by the vault owner
          - waiting_on: [{person, description}] — owed TO the owner
          - thread_updates: [{slug, body}] — existing-thread evidence
          - new_threads: [{slug, thesis}] — brand-new threads
          - concepts: [str] — concept slugs

        Notes never have attendees, decisions, or an explicit date.
        """
        return _do_structured_record(
            config=config, source_type="note",
            title=title, body=body, date=None, attendees=[],
            result_payload={
                "summary": summary,
                "attendees": [],
                "decisions": [],
                "action_items_mine": action_items_mine or [],
                "waiting_on": waiting_on or [],
                "other_commitments": [],
                "thread_updates": thread_updates or [],
                "new_threads": new_threads or [],
                "concepts": concepts or [],
            },
        )

    @mcp.tool()
    def record_doc(
        title: str,
        body: str,
        summary: str,
        attendees: list[Attendee] | None = None,
        action_items_mine: list[str] | None = None,
        waiting_on: list[PersonItem] | None = None,
        thread_updates: list[ThreadUpdate] | None = None,
        new_threads: list[NewThread] | None = None,
        concepts: list[str] | None = None,
    ) -> dict:
        """Persist a doc / strategy brief / slide deck / competitive report.

        Nested types (see schema in the tool's inputSchema):
          - attendees: [{name, role?, email?}] — pass [] for docs with no
            named author or participant list; don't invent authors.
          - action_items_mine: [str] — items owned by the vault owner;
            pass [] if the doc doesn't assign work.
          - waiting_on: [{person, description}] — owed TO the owner
          - thread_updates: [{slug, body}] — existing-thread evidence.
            NOT optional-in-spirit for docs: if the doc connects to any
            existing thread, this is how it shows up on those threads.
          - new_threads: [{slug, thesis}] — brand-new threads to track
          - concepts: [str] — concept slugs; almost always worth populating

        The doc still gets indexed, summarized, and searchable even with
        empty attendees / action items. thread_updates, new_threads, and
        concepts are how docs connect to the rest of the vault.
        """
        return _do_structured_record(
            config=config, source_type="doc",
            title=title, body=body, date=None,
            attendees=attendees or [],
            result_payload={
                "summary": summary,
                "attendees": attendees or [],
                "decisions": [],
                "action_items_mine": action_items_mine or [],
                "waiting_on": waiting_on or [],
                "other_commitments": [],
                "thread_updates": thread_updates or [],
                "new_threads": new_threads or [],
                "concepts": concepts or [],
            },
        )

    # ---------- LEGACY: LLM-backed ingest (requires ANTHROPIC_API_KEY) ----------
    #
    # Kept so a user running their own API key can do batch ingest from the
    # CLI or tools outside Claude Desktop. For the default chat workflow,
    # use the structured record_* tools above instead — those don't need an
    # API key because Claude Desktop does the analysis.

    def _require_api_key() -> dict | None:
        import os
        if not os.environ.get("ANTHROPIC_API_KEY"):
            return {
                "error": (
                    "This tool runs its own LLM call and requires "
                    "ANTHROPIC_API_KEY to be set on the server. For the "
                    "no-API-key workflow, use `record_meeting` / "
                    "`record_note` / `record_doc` (Claude Desktop does the "
                    "analysis)."
                )
            }
        return None

    @mcp.tool()
    def ingest_note(text: str, title: str | None = None) -> dict:
        """LEGACY: file a note with server-side LLM analysis.

        Requires ANTHROPIC_API_KEY. Prefer `record_note` for the no-API-key
        flow (Claude Desktop analyzes, this server just persists).
        """
        gate = _require_api_key()
        if gate:
            return gate
        safe_title = (title or f"note-{datetime.now().strftime('%Y%m%d-%H%M%S')}").strip()
        slug = _slugify(safe_title)
        dest = config.raw_notes / f"{slug}.md"
        config.raw_notes.mkdir(parents=True, exist_ok=True)
        dest.write_text(text)
        _read_new_source(config, dest, "note")
        return {"ingested": str(dest), "source_type": "note"}

    @mcp.tool()
    def ingest_meeting(
        text: str,
        title: str | None = None,
        date_str: str | None = None,
        attendees: list[str] | None = None,
    ) -> dict:
        """LEGACY: file a meeting with server-side LLM analysis.

        Requires ANTHROPIC_API_KEY. Prefer `record_meeting` for the
        no-API-key flow.
        """
        gate = _require_api_key()
        if gate:
            return gate
        from deep_reader.sources.meeting import parse_meeting
        from deep_reader.markdown import format_frontmatter
        meta = parse_meeting(text)
        the_title = title or meta.title or "Meeting"
        the_date = date_str or (meta.meeting_date.isoformat() if meta.meeting_date else None)
        the_attendees = attendees if attendees is not None else meta.attendees
        fm: dict = {"title": the_title, "type": "meeting"}
        if the_date:
            fm["date"] = the_date
        if the_attendees:
            fm["attendees"] = the_attendees
        slug = _slugify(the_title)
        filename = f"{the_date}-{slug}.md" if the_date else f"{slug}.md"
        config.raw_meetings.mkdir(parents=True, exist_ok=True)
        dest = config.raw_meetings / filename
        dest.write_text(format_frontmatter(fm) + text)
        _read_new_source(config, dest, "meeting")
        return {"ingested": str(dest), "source_type": "meeting"}

    @mcp.tool()
    def recap_prep(date_str: str | None = None) -> dict:
        """Write a prep file for the daily-recap skill. Returns the file path."""
        sys.path.insert(0, str(Path(__file__).parent.parent / "tools"))
        from recap_prep import run_recap_prep
        td = date.fromisoformat(date_str) if date_str else None
        path = run_recap_prep(config, td)
        return {"path": str(path), "content": path.read_text()}

    @mcp.tool()
    def sync_recap(date_str: str | None = None) -> dict:
        """Pull action items from a daily-recap file into the wiki."""
        sys.path.insert(0, str(Path(__file__).parent.parent / "tools"))
        from sync_recap import run_sync_recap
        td = date.fromisoformat(date_str) if date_str else None
        added = run_sync_recap(config, td)
        return {"added": added}

    # ---------- Prompts ----------
    #
    # Prompts are saved workflows surfaced by Claude Desktop. In this
    # architecture Claude does the reading, parsing, and summarization
    # itself, then calls the structured `record_*` tools to persist. The
    # server never makes LLM calls in this flow — Claude's own subscription
    # covers all the model work.

    ANALYZE_SCHEMA = """
Pick the right record_* tool for the source type:
- record_meeting — conversations with attendees + decisions + follow-ups
- record_doc — docs, strategy briefs, slide decks, competitive briefs,
  research reports, one-pagers, anything authored rather than discussed.
  Most of these have no attendees and no action items — that's fine.
- record_note — short personal note, clip, jotting

Fields (ALL are optional except title/body/summary — pass only what's
present in the source, omit or pass [] for anything that isn't):
- title (string, required)
- body (string, required) — the original text, unmodified
- summary (string, required) — 2-4 sentences capturing the gist
- date (YYYY-MM-DD, meetings only, if known)
- attendees (list) — each {name, role?, email?}. Omit for docs/decks
  with no explicit author or participant list. For docs WITH authors,
  pass the author as a single-entry attendees list with their role.
- decisions (list of strings) — meetings only, usually
- action_items_mine (list of strings) — items owned by the vault owner
  ONLY. Check `get_ingest_context().owner` for who that is; items owned
  by anyone else belong in waiting_on or other_commitments, NOT here.
  Empty [] is common for docs that don't assign work.
- waiting_on (list) — each {person, description}. Items owed TO the
  vault owner by a named person. Empty [] is fine.
- other_commitments (list) — each {person, description}. Items between
  other parties that don't involve the vault owner.
- thread_updates (list) — each {slug, body}. For every thread in
  get_ingest_context().threads whose thesis this source meaningfully
  advances, produce a one-sentence evidence entry (not a rewrite of the
  thesis). THIS MATTERS EVEN FOR DOCS WITH NO PEOPLE — a competitive
  brief can and should advance threads it relates to.
- new_threads (list) — each {slug, thesis}. Only for genuinely recurring
  themes introduced here that would be worth tracking across future sources.
- concepts (list of strings) — concept slugs relevant to this source.
  Important even for docs — how you graduate ideas to concept articles.
""".strip()

    @mcp.prompt(
        description=(
            "Analyze a pasted meeting note and record it in the vault. "
            "I'll paste the meeting content as my next message."
        ),
    )
    def ingest_meeting_paste() -> str:
        return (
            "I'll paste a meeting note as my next message. Do these steps in "
            "order:\n"
            "1. Call `get_ingest_context()` so you know the vault owner, "
            "active threads (with theses), and existing people.\n"
            "2. Read the pasted content carefully.\n"
            "3. Call `record_meeting(...)` with the structured analysis.\n"
            "\n"
            f"{ANALYZE_SCHEMA}\n"
            "\n"
            "Critical: before finalizing action_items_mine, compare each item "
            "against the vault owner's name/aliases from get_ingest_context. "
            "An item is 'mine' only if the OWNER is doing it. 'I'll send the "
            "deck' said by the owner → mine. 'Jane will send the deck' → "
            "waiting_on with person='Jane'.\n"
            "\n"
            "After recording, summarize what changed: attendees added, new "
            "people, new action items on my list, threads advanced."
        )

    @mcp.prompt(
        description=(
            "Analyze a pasted doc, brief, slide deck, or report and record "
            "it in the vault. Works for sources with no attendees / no "
            "action items — still pulls out threads and concepts."
        ),
    )
    def ingest_doc_paste() -> str:
        return (
            "I'll paste a doc, brief, slide deck, or report as my next "
            "message. Steps:\n"
            "1. Call `get_ingest_context()` for threads, people, owner.\n"
            "2. Read the content.\n"
            "3. Call `record_doc(...)` with the structured analysis. If "
            "the doc has no explicit author/attendee list, pass "
            "attendees=[] and don't invent anyone. If it has no action "
            "items assigned to anyone, pass action_items_mine=[] and "
            "waiting_on=[]. That's normal for a strategy doc or brief.\n"
            "\n"
            "The fields that MUST be populated even for a people-less doc: "
            "summary, thread_updates (if it meaningfully advances any "
            "existing thread), new_threads (if it introduces a recurring "
            "theme), concepts (always). The whole point is this doc "
            "still connects to existing work even without people attached.\n"
            "\n"
            f"{ANALYZE_SCHEMA}\n"
            "\n"
            "After recording, summarize: threads advanced or created, "
            "concepts tagged, anything that now connects to existing "
            "work in the vault."
        )

    @mcp.prompt(
        description=(
            "Process every file in the inbox — read each, analyze it, and "
            "record it in the vault. Skips anything already present."
        ),
    )
    def ingest_inbox() -> str:
        return (
            "Process the vault inbox in order:\n"
            "1. Call `list_inbox()`.\n"
            "2. Call `get_ingest_context()` once — reuse it across every "
            "file so your 'mine' classification stays consistent.\n"
            "3. For each file:\n"
            "   a. Call `read_inbox_file(filename)` to get the content and "
            "suggested type.\n"
            "   b. **Dedup check.** Derive a likely title from the content "
            "or filename, then call `search(query=<title or distinctive "
            "phrase>)`. If the first result's source slug looks like a "
            "match (same title, same date), skip this file — call "
            "`move_inbox_file(filename, source_type)` with the existing "
            "source type and report it as 'already in vault'. Do NOT call "
            "record_* for files you believe are duplicates.\n"
            "   c. For new files: pick the right record_* tool "
            "(record_meeting / record_note / record_doc) based on content "
            "and the suggested_type hint from read_inbox_file.\n"
            "   d. Call that tool with your structured analysis.\n"
            "   e. On success, call `move_inbox_file(filename, "
            "source_type)` to archive the original.\n"
            "4. At the end, give me a summary — how many were new, how "
            "many skipped as duplicates, any that failed and why, and the "
            "top new action items on my list.\n"
            "\n"
            f"{ANALYZE_SCHEMA}"
        )

    @mcp.prompt(
        description=(
            "Fetch today's meetings from Granola and record each in the "
            "vault. Requires Granola's MCP server to also be registered."
        ),
    )
    def ingest_granola_today() -> str:
        return (
            "Pull today's Granola meetings and record them:\n"
            "1. Call `get_ingest_context()` for owner, threads, people.\n"
            "2. Use the Granola MCP tools to list every meeting from today.\n"
            "3. For each meeting:\n"
            "   a. Get its full content, title, date, and attendee list "
            "from Granola's tools.\n"
            "   b. Analyze it against the ingest context.\n"
            "   c. Call `record_meeting(...)` with the structured result.\n"
            "4. Summarize what was added.\n"
            "If Granola returns nothing for today, say so — don't invent.\n"
            "\n"
            f"{ANALYZE_SCHEMA}"
        )

    @mcp.prompt(
        description=(
            "Fetch Granola meetings for a date range and record each. Args: "
            "start_date, end_date (YYYY-MM-DD)."
        ),
    )
    def ingest_granola_range(start_date: str, end_date: str) -> str:
        return (
            f"Pull Granola meetings between {start_date} and {end_date} "
            f"(inclusive) and record them:\n"
            "1. Call `get_ingest_context()`.\n"
            "2. Use Granola's tools to list meetings in the range.\n"
            "3. For each, check if it's already in the vault (call "
            "`search(query=meeting_title)`). Skip any that match.\n"
            "4. For new ones, analyze and call `record_meeting(...)`.\n"
            "5. Summarize additions.\n"
            "\n"
            f"{ANALYZE_SCHEMA}"
        )

    @mcp.prompt(
        description=(
            "Weekly catchup — fetch last 7 days of Granola meetings and "
            "record any not already in the vault."
        ),
    )
    def ingest_granola_week() -> str:
        return (
            "Run a weekly catchup from Granola:\n"
            "1. Call `get_ingest_context()`.\n"
            "2. List Granola meetings from the past 7 days.\n"
            "3. For each, call `search(query=meeting_title)` to check for "
            "duplicates. Skip matches.\n"
            "4. For new meetings, analyze and `record_meeting(...)`.\n"
            "5. Give me a weekly digest: who I met with most, recurring "
            "themes (reference any threads that advanced), top new items.\n"
            "\n"
            f"{ANALYZE_SCHEMA}"
        )

    @mcp.prompt(
        description=(
            "Pull today's notes from your personal Slack channel and "
            "ingest them — captures self-todos and reminders, files the "
            "day as a note source. Dedupes against existing items to "
            "avoid double-counting reaffirmations. Requires the Slack "
            "MCP server registered alongside this one."
        ),
    )
    def ingest_slack_personal(date_str: str = "today") -> str:
        return (
            f"Pull and ingest my Slack personal-channel notes from "
            f"{date_str}. Steps:\n"
            "\n"
            "1. Call `get_ingest_context()` for owner identity and "
            "active threads.\n"
            "2. Call `list_action_items(status='open')` — hold this "
            "list; you'll dedup against it in step 6.\n"
            "3. Use the Slack MCP to find my personal channel. It's "
            "typically: a DM I have with myself, OR a private channel "
            "I use for self-notes (often named like #my-notes / "
            "#scratch / #dan-todo or similar). If you can't tell which "
            "channel is the personal one, ASK me before guessing.\n"
            "4. Read messages from that channel for the target date.\n"
            "5. Check for duplicate-day ingest: call `search(query="
            f"'slack personal {date_str}', depth='lite')`. If a source "
            "like 'note-slack-personal-<date>' already exists, ask "
            "whether to skip, append, or replace before proceeding.\n"
            "6. **Pre-classify candidate todos before record_note**: "
            "for each actionable item you'd extract from the messages, "
            "compare against the existing-items list from step 2. "
            "Categorize each as:\n"
            "   - **NEW**: not yet on my list — include in "
            "`action_items_mine` for the record_note call.\n"
            "   - **PARAPHRASE of existing item**: do NOT include in "
            "action_items_mine. After record_note completes, call "
            "`link_action_item(id=<existing-id>, source_ref=<Slack "
            "permalink>)` for each one to attach the Slack reference "
            "to the existing item.\n"
            "7. Bundle the day's content as a single note via "
            "`record_note(...)`:\n"
            "   - title: 'Slack personal notes — <YYYY-MM-DD>'\n"
            "   - body: the messages verbatim, with timestamps and "
            "permalinks if available\n"
            "   - summary: 1-3 sentences on the day's themes\n"
            "   - action_items_mine: only the NEW items from step 6\n"
            "   - waiting_on: [{person, description}] for things I "
            "noted I'm owed by named others\n"
            "   - thread_updates: only if the day's notes meaningfully "
            "advance an existing thread\n"
            "   - concepts: tag recurring themes\n"
            "8. After record_note: call `link_action_item` for each "
            "PARAPHRASE identified in step 6.\n"
            "9. Report what landed: NEW items added, PARAPHRASES "
            "linked to existing, threads advanced, duplicates skipped.\n"
            "\n"
            "Bias toward linking, not adding. Personal channels often "
            "re-mention things already on the list ('still need to send "
            "deck') — those should attach to the existing item, not "
            "duplicate it. Half-formed thoughts with no clear next step: "
            "leave in body, don't extract.\n"
        )

    @mcp.prompt(
        description=(
            "Scan today's Slack chats across channels for action items "
            "you committed to or that others owe you. Extracts items "
            "into your central list, deduping against items already "
            "captured from meetings or earlier ingests. Does NOT create "
            "source pages for regular chat — use ingest_slack_thread "
            "for substantive threads worth their own page."
        ),
    )
    def ingest_slack_action_items(date_str: str = "today") -> str:
        return (
            f"Scan my Slack chats from {date_str} for action items "
            f"only — no source pages, just extract commitments. The "
            f"key skill here is dedup against items already captured "
            f"from meetings (a Slack reaffirmation of a commitment is "
            f"common and shouldn't double up).\n"
            "\n"
            "Steps:\n"
            "1. Call `get_ingest_context()` for owner identity and "
            "known people.\n"
            "2. Call `list_action_items(status='open')` and "
            "`list_waiting_on(status='open')`. Hold this list — you'll "
            "compare every Slack candidate against it.\n"
            "3. Use the Slack MCP to identify channels I was active in "
            "for the target date. DMs, small group chats, and project "
            "channels matter; large broadcast channels usually don't.\n"
            "4. For each relevant channel/thread, look for action-item-"
            "shaped commitments:\n"
            "   - Things I said I'd do ('I'll send the deck')\n"
            "   - Things someone else said they'd do for me ('Jane will "
            "follow up by Friday')\n"
            "   - Explicit asks where the answer was a clear yes\n"
            "5. **Dedup decision per candidate**, against the list from "
            "step 2:\n"
            "   - **Exact paraphrase of an existing item** (same intent, "
            "different wording — e.g., 'send the deck Friday' vs "
            "existing 'Send pricing deck to Jane by Friday'): call "
            "`link_action_item(id=<existing-id>, source_ref=<Slack "
            "permalink>)`. Do NOT call add_action_item.\n"
            "   - **New item**, owned by me: call `add_action_item("
            "description, source=<Slack permalink or 'slack:#channel'>)`\n"
            "   - **New item**, owed to me by named person: call "
            "`add_waiting_on(description, person, source=<permalink>)`\n"
            "   - **Aspirational / speculative** ('we should probably do "
            "X'): skip.\n"
            "   - **Between other parties / no clear owner**: skip.\n"
            "6. Report: count of new items added, count of links to "
            "existing items, channels scanned, anything skipped.\n"
            "\n"
            "Bias toward linking, not adding. If you're unsure whether "
            "a Slack commitment is the same thing as an existing item, "
            "lean toward linking — it's lossless (you can always split "
            "later). Adding a duplicate item is harder to undo."
        )

    @mcp.prompt(
        description=(
            "Ingest a substantive Slack thread as its own source — for "
            "when a back-and-forth conversation produced enough decision "
            "/ context to warrant being its own page, not just an "
            "action item. You'll provide the channel + thread."
        ),
    )
    def ingest_slack_thread() -> str:
        return (
            "I'll point you at a specific Slack channel and thread to "
            "ingest as a source. Steps:\n"
            "\n"
            "1. Call `get_ingest_context()`.\n"
            "2. Use the Slack MCP to read the full thread (parent "
            "message + all replies).\n"
            "3. Treat this as a meeting analog — it has participants, "
            "decisions, action items. Call `record_meeting(...)`:\n"
            "   - title: a short descriptive title for the conversation "
            "(infer from the parent message)\n"
            "   - date: thread date\n"
            "   - body: full thread text with usernames and timestamps\n"
            "   - attendees: everyone who posted, with their names "
            "(resolve via Slack user lookup if needed). Pass the vault "
            "owner as one of the attendees.\n"
            "   - summary, decisions, action_items_mine, waiting_on, "
            "thread_updates, new_threads, concepts — same as a regular "
            "meeting ingest.\n"
            "4. Report what landed.\n"
            "\n"
            "If the thread is light (a few messages, no real decisions), "
            "use `record_note(...)` instead of record_meeting — same "
            "fields minus attendees/decisions/date.\n"
        )

    # ---------- Concept distillation prompts (the synthesis exception) ----------
    #
    # Concept pages are the ONE place we synthesize. They're meta-entities
    # that exist as integrations across sources. All operations go through
    # the review queue so the user approves before pages are written.

    @mcp.prompt(
        description=(
            "Refresh a concept page — re-reads the sources tagging this "
            "concept and proposes an updated definition + distillation. "
            "Goes to the review queue so you can diff before it's written."
        ),
    )
    def refresh_concept(name: str) -> str:
        return (
            f"Refresh the concept page for '{name}'. The result goes to "
            f"the review queue for approval — does NOT immediately write "
            f"the page.\n"
            "\n"
            f"1. `get_concept_evidence(name='{name}')` — every source tagging "
            f"this concept, with full content. Also returns existing_article "
            f"if there's already a page.\n"
            f"2. `get_concept_with_hierarchy(name='{name}')` — current "
            f"parent/child/related concepts.\n"
            f"3. Read every source's content. Identify how the concept is "
            f"used / understood across sources. Note tensions if sources "
            f"disagree.\n"
            f"4. Write a structured concept page with sections:\n"
            f"   - **Definition** (1 paragraph): how the concept is used in "
            f"this vault's context. NOT a generic dictionary def.\n"
            f"   - **Distillation** (2-3 paragraphs): integrated "
            f"understanding. Heavy citation via `[[<source-slug>]]`. Use "
            f"direct markdown blockquotes where source language matters. "
            f"This is the place where it's OK to synthesize across sources, "
            f"BUT every claim must be traceable.\n"
            f"   - **Tensions / open questions** (optional): where sources "
            f"disagree or the concept is still evolving, with citations.\n"
            f"5. Call `propose_review(\n"
            f"      kind='concept_refresh',\n"
            f"      title='Refresh concept page: {name}',\n"
            f"      preview=<diff-style summary: what's changing vs existing page (if any)>,\n"
            f"      proposed_action={{\n"
            f"          'tool': 'record_concept_page',\n"
            f"          'args': {{name, definition, distillation, "
            f"contributing_sources, parent_concepts, child_concepts, "
            f"related_concepts, tensions}}\n"
            f"      }}\n"
            f"   )`.\n"
            f"6. Tell me the review id and the title — say *'review {{id}} "
            f"is queued. Approve in chat with: approve {{id}}'*.\n"
            "\n"
            "Hard rule: if you can't write the distillation without "
            "paraphrasing every source into abstract prose, say the "
            "evidence is too thin and skip — don't queue a review."
        )

    @mcp.prompt(
        description=(
            "Survey vault for stale concept pages and propose refreshes. "
            "Each refresh goes to the review queue."
        ),
    )
    def list_stale() -> str:
        return (
            "Survey the vault for stale concept pages and propose refreshes:\n"
            "\n"
            "1. `list_stale_concepts(min_new_sources=3)`.\n"
            "2. For each (or batch — ask if many): run the refresh_concept "
            "flow and queue a review item. Don't actually write anything; "
            "the user approves via the queue.\n"
            "3. Report the list of queued reviews with IDs."
        )

    @mcp.prompt(
        description=(
            "Survey ingested sources, propose concept hierarchy "
            "relationships (parent/child/related). Goes to review queue."
        ),
    )
    def suggest_concept_links() -> str:
        return (
            "Survey the vault for concept hierarchy you might be missing.\n"
            "\n"
            "1. `list_concept_candidates(min_sources=2)` — get every "
            "concept currently in the vault.\n"
            "2. For each pair of concepts, look at the sources that tag "
            "them. If sources strongly suggest 'A is a kind of B' or 'A "
            "and B are aspects of the same thing', propose a relationship.\n"
            "3. For each proposal, call `propose_review(\n"
            "     kind='concept_link',\n"
            "     title='Link: <child> → <parent>',\n"
            "     preview=<one-paragraph rationale citing the sources that "
            "suggested it>,\n"
            "     proposed_action={'tool': 'link_concepts', 'args': "
            "{parent, child, kind}}\n"
            "   )`.\n"
            "4. Report all queued reviews to me with their IDs.\n"
            "\n"
            "Conservative bias: don't propose if you wouldn't be confident "
            "saying it out loud. False positives are noise."
        )

    # ---------- Drive integration prompts ----------

    @mcp.prompt(
        description=(
            "One-time backfill from Google Drive — walks every doc in a "
            "folder, ingests substantive content. Heavy operation; run "
            "per-folder. Requires Drive MCP server registered alongside."
        ),
    )
    def backfill_drive(folder_path: str = "") -> str:
        target = folder_path or "(ask me which folder to start with)"
        return (
            f"Backfill Drive content from {target}:\n"
            "\n"
            "1. Use the Drive MCP to list every doc in the target folder "
            "(recursive). Filter to recognizable content types: Google "
            "Docs, PDFs, Slides, .docx files. Skip Sheets unless they're "
            "structured docs, skip forms / junk.\n"
            "2. Call `list_drive_ingested()` to get already-ingested doc "
            "IDs. Skip those.\n"
            "3. Call `get_ingest_context()` once for owner identity, "
            "active threads, known people.\n"
            "4. For each remaining doc:\n"
            "   a. Read full content via the Drive MCP (don't rely on "
            "metadata alone — actual content determines whether it's "
            "substantive).\n"
            "   b. If the doc is genuinely junk (auto-generated forms, "
            "expired drafts, random notes with no structure), skip and "
            "report why.\n"
            "   c. Otherwise: analyze and call `record_doc(...)` with "
            "the structured payload.\n"
            "   d. After successful record_doc, call "
            "`mark_drive_ingested(drive_id, source_slug)` to track it.\n"
            "5. Report: total docs found, ingested, skipped (with "
            "reasons), failed (with errors).\n"
            "\n"
            "This is heavy — could be 50–500 docs per folder. Do them "
            "in batches, ask before continuing if the folder is huge. "
            "If a doc's relevance is borderline, instead of skipping or "
            "ingesting, call `propose_review(kind='drive_borderline', "
            "title='Borderline: <doc title>', preview=<one-paragraph "
            "summary + rationale>, proposed_action={'tool': "
            "'ingest_drive_doc', 'args': {drive_id, content, ...}})` and "
            "let me decide later."
        )

    @mcp.prompt(
        description=(
            "Incremental Drive crawl — process docs modified since last "
            "crawl. Cheaper than backfill; run weekly or on demand."
        ),
    )
    def crawl_drive(since: str = "") -> str:
        return (
            f"Incremental Drive crawl. Only docs modified since "
            f"{since or 'last crawl'}.\n"
            "\n"
            "1. `list_drive_ingested()` — get last_crawl_at and ingested IDs.\n"
            "2. Use Drive MCP to list recently-modified docs since the "
            "boundary (parameter `since` if provided, else last_crawl_at, "
            "else 7 days ago).\n"
            "3. `get_ingest_context()` once.\n"
            "4. For each doc, same flow as backfill_drive: skip if "
            "already ingested, read content, analyze + record_doc + "
            "mark_drive_ingested. Borderline cases → review queue.\n"
            "5. Report what was processed."
        )

    # ---------- Proactive enrichment prompts ----------

    @mcp.prompt(
        description=(
            "Find Drive/Linear material related to a concept and propose "
            "ingesting it. Each candidate goes to the review queue."
        ),
    )
    def enrich_concept(name: str) -> str:
        return (
            f"Proactively find material in Drive (and Linear if registered) "
            f"that should enrich the '{name}' concept.\n"
            "\n"
            f"1. `get_concept_with_hierarchy(name='{name}')` and "
            f"`get_concept_evidence(name='{name}')` — current scope.\n"
            "2. If a Drive MCP is registered: search Drive (full-text, "
            f"not just metadata) for content related to '{name}' and its "
            "hierarchy keywords. Filter out things already in the vault "
            "via `is_drive_ingested(drive_id)`.\n"
            "3. If a Linear MCP is registered: search Linear for issues "
            "tagged or mentioning the concept.\n"
            "4. For each promising candidate, call `propose_review(\n"
            "     kind='enrichment_ingest',\n"
            f"     title='Ingest for {name}: <doc/issue title>',\n"
            "     preview=<one-paragraph why this is relevant + a 200-char "
            "excerpt>,\n"
            "     proposed_action={'tool': 'ingest_drive_doc' or 'ingest_linear_issue', "
            "'args': <enough metadata for the actual ingest>}\n"
            "   )`.\n"
            "5. Report: how many candidates queued, with IDs.\n"
            "\n"
            "Conservative bias — only propose if the candidate would "
            "obviously enrich the concept. Don't pad."
        )

    @mcp.prompt(
        description=(
            "Find Drive/Linear material related to a thread and propose "
            "ingesting it. Each candidate goes to the review queue."
        ),
    )
    def enrich_thread(slug: str) -> str:
        return (
            f"Proactively find Drive/Linear material that should enrich "
            f"the '{slug}' thread.\n"
            "\n"
            f"1. `get_thread_full_context(slug='{slug}')` — current scope.\n"
            "2. Use the thread thesis + recent evidence as a search query "
            "against Drive (and Linear if registered). Skip already-"
            "ingested docs.\n"
            "3. For each candidate, propose_review with kind="
            "'enrichment_ingest', title naming the thread, preview "
            "explaining the relevance.\n"
            "4. Report queued reviews."
        )

    @mcp.prompt(
        description=(
            "Find Drive/Linear material related to a person and propose "
            "ingesting it. Each candidate goes to the review queue."
        ),
    )
    def enrich_person(name: str) -> str:
        return (
            f"Proactively find Drive/Linear material involving {name}.\n"
            "\n"
            f"1. `get_person_full_context(slug='{name}')`.\n"
            "2. Search Drive for docs they authored, were shared with, "
            "or are mentioned in. Use their email + name + aliases.\n"
            "3. Search Linear for issues they own / are assigned / "
            "mentioned in.\n"
            "4. For each new candidate, propose_review with kind="
            "'enrichment_ingest'.\n"
            "5. Report queued reviews."
        )

    # ---------- Review queue interaction prompt ----------

    @mcp.prompt(
        description=(
            "Walk through pending review items — show each, let me approve "
            "or reject in chat."
        ),
    )
    def review_pending() -> str:
        return (
            "Walk me through pending reviews:\n"
            "\n"
            "1. `list_pending_reviews()`.\n"
            "2. Group by kind. For each kind, show the count.\n"
            "3. For each item: show its title, id, and preview. Ask "
            "approve/reject/skip.\n"
            "4. On approve: call `approve_review(id)` — the proposed "
            "action executes.\n"
            "5. On reject: call `reject_review(id, reason=<my reason>)`.\n"
            "6. On skip: leave it queued, move to next.\n"
            "7. At the end, summarize: how many approved, rejected, "
            "still pending.\n"
            "\n"
            "Show me ONE item at a time unless I tell you to batch."
        )

    @mcp.prompt(
        description=(
            "Catch me up — brief on what's changed in the vault: open "
            "items, new people, recent sources, pending reviews."
        ),
    )
    def catch_me_up() -> str:
        return (
            "Give me a short brief (~200 words max). Steps:\n"
            "1. Read the `vault://summary` resource.\n"
            "2. Call `list_action_items(status='open')` and "
            "`list_waiting_on(status='open')`.\n"
            "3. Report: my top open items ranked by age, anything I'm "
            "waiting on that's been sitting too long, the 3-5 most recent "
            "sources, and one concrete thing to do next."
        )

    @mcp.prompt(
        description=(
            "Quick scan — list what the vault has on a term, no synthesis. "
            "Use when you just want to see WHAT sources/threads mention "
            "something, not a detailed answer. Slash-command shortcut for "
            "the lightweight search path."
        ),
    )
    def quick_scan(term: str) -> str:
        return (
            f"Do a quick lightweight scan for '{term}' and return just a "
            f"tight list — no synthesis, no reasoning beyond what's in the "
            f"vault.\n"
            "\n"
            f"Steps:\n"
            f"1. Call `search(query='{term}', depth='lite')`.\n"
            "2. Format a short bullet list of what turned up, grouped by "
            "type (sources, threads, concepts, people, action items). "
            "For each hit show just the name/slug and a one-line snippet.\n"
            "3. If nothing matched, say so — don't pad.\n"
            "\n"
            "Do NOT fetch full content, do NOT try to answer a question "
            "from these results. This is just a scan."
        )

    @mcp.prompt(
        description=(
            "Force deep retrieval for a substantive question. Usually "
            "unnecessary — default `search` already returns full content "
            "of top hits. Keep this as a fallback if Claude ever answers "
            "from snippets instead of retrieved content."
        ),
    )
    def deep_query(question: str) -> str:
        return (
            f"Answer this question by reading the vault properly, not by "
            f"reasoning from search snippets: {question}\n"
            "\n"
            "Follow this exact pattern:\n"
            "1. Call `search(query=...)` with a targeted query. Note "
            "which sources and threads scored highest.\n"
            "2. For each top source hit that looks relevant, call "
            "`get_source(slug)` to read its full content (overview + "
            "chunk pages with decisions, action items, concepts).\n"
            "3. For each top thread hit, fetch `vault://threads/{name}` "
            "to see the full evidence log, not just the thesis.\n"
            "4. If a person is central to the question, call "
            "`get_person(name)` for their full interaction history.\n"
            "5. ONLY THEN synthesize the answer. Cite specific sources "
            "and quote actual content when the question is about what "
            "was said or decided.\n"
            "6. If after retrieval you still don't have enough to answer "
            "well, say what's missing rather than making it up.\n"
            "\n"
            "The goal is an answer grounded in the vault, not an answer "
            "reconstructed from your general knowledge of the topic. If "
            "the vault has the substance, your answer should quote or "
            "paraphrase the vault, not restate what's generally true."
        )

    return mcp


# ---------- helpers ----------


def _entity_sources(state, slug: str, config: Config) -> set:
    """Return the set of source slugs associated with any entity slug.

    Source: just {slug} itself. Person: their appearances. Thread:
    sources in evidence. Concept: sources tagging it.
    """
    if slug in state.sources:
        return {slug}
    if slug in state.people:
        return set(state.people[slug].appearances)
    if slug in state.global_threads:
        thread_path = config.wiki_threads / f"{slug}.md"
        if thread_path.exists():
            from deep_reader.thread_utils import extract_section
            import re as _re
            evidence = extract_section(thread_path.read_text(), "Evidence")
            return {m.group(1) for m in _re.finditer(r"\[\[([^\]/]+)/chunk-\d+\]\]", evidence)}
    if slug in state.concepts or slug.lower() in state.concepts:
        slug_l = slug.lower() if slug.lower() in state.concepts else slug
        import re as _re
        out: set = set()
        for src_slug in state.sources:
            src_dir = config.wiki_sources / src_slug
            for cp in src_dir.glob("chunk-*.md"):
                if _re.search(r"\[\[" + _re.escape(slug_l) + r"\]\]", cp.read_text(), _re.IGNORECASE):
                    out.add(src_slug)
                    break
        return out
    # Try as person name
    for p in state.people.values():
        if p.name.lower() == slug.lower():
            return set(p.appearances)
    return set()


def _related_for_source(state, slug: str, limit: int) -> dict:
    """Co-attendees, threads advanced, concepts tagged."""
    src = state.sources[slug]
    co_attendees = []
    for p in state.people.values():
        if slug in p.appearances:
            co_attendees.append({"slug": p.slug, "name": p.name})
    return {
        "kind": "source",
        "slug": slug,
        "co_attendees": co_attendees[:limit],
        "threads_in_source": list(src.threads)[:limit],
    }


def _related_for_person(state, slug: str, limit: int) -> dict:
    """Sources, threads, frequent co-appearances."""
    p = state.people[slug]
    # Co-appearance ranking: how many sources do other people share
    co_app: dict = {}
    p_set = set(p.appearances)
    for other in state.people.values():
        if other.slug == slug:
            continue
        shared = len(p_set & set(other.appearances))
        if shared:
            co_app[other.slug] = {"slug": other.slug, "name": other.name, "shared": shared}
    co_app_sorted = sorted(co_app.values(), key=lambda x: x["shared"], reverse=True)
    return {
        "kind": "person",
        "slug": slug,
        "name": p.name,
        "appearances": p.appearances[-limit:],
        "frequent_co_appearances": co_app_sorted[:limit],
    }


def _related_for_thread(state, slug: str, config: Config, limit: int) -> dict:
    """Contributing sources, central people, related threads."""
    thread_path = config.wiki_threads / f"{slug}.md"
    contributing: list = []
    if thread_path.exists():
        from deep_reader.thread_utils import extract_section
        import re as _re
        evidence = extract_section(thread_path.read_text(), "Evidence")
        seen: set = set()
        for m in _re.finditer(r"\[\[([^\]/]+)/chunk-\d+\]\]", evidence):
            if m.group(1) not in seen:
                seen.add(m.group(1))
                contributing.append(m.group(1))
    contrib_set = set(contributing)
    # Central people (most appearances in contributing)
    people_overlap = []
    for p in state.people.values():
        c = len(set(p.appearances) & contrib_set)
        if c:
            people_overlap.append({"slug": p.slug, "name": p.name, "in_count": c})
    people_overlap.sort(key=lambda x: x["in_count"], reverse=True)
    # Related threads: other threads sharing 2+ sources
    related_threads = []
    for other_slug in state.global_threads:
        if other_slug == slug:
            continue
        other_path = config.wiki_threads / f"{other_slug}.md"
        if not other_path.exists():
            continue
        from deep_reader.thread_utils import extract_section
        import re as _re
        other_ev = extract_section(other_path.read_text(), "Evidence")
        other_set = {m.group(1) for m in _re.finditer(r"\[\[([^\]/]+)/chunk-\d+\]\]", other_ev)}
        shared = len(contrib_set & other_set)
        if shared >= 2:
            related_threads.append({"slug": other_slug, "shared_sources": shared})
    related_threads.sort(key=lambda x: x["shared_sources"], reverse=True)
    return {
        "kind": "thread",
        "slug": slug,
        "contributing_sources": contributing[:limit],
        "central_people": people_overlap[:limit],
        "related_threads": related_threads[:limit],
    }


def _related_for_concept(state, slug: str, config: Config, limit: int) -> dict:
    """Contributing sources + hierarchy + central people."""
    c = state.concepts.get(slug)
    if not c:
        return {"error": f"concept '{slug}' not found"}
    # Contributing sources via concept tags
    import re as _re
    contributing: list = []
    for src_slug in state.sources:
        src_dir = config.wiki_sources / src_slug
        for cp in src_dir.glob("chunk-*.md"):
            if _re.search(r"\[\[" + _re.escape(slug) + r"\]\]", cp.read_text(), _re.IGNORECASE):
                contributing.append(src_slug)
                break
    contrib_set = set(contributing)
    # Central people
    people_overlap = []
    for p in state.people.values():
        cnt = len(set(p.appearances) & contrib_set)
        if cnt:
            people_overlap.append({"slug": p.slug, "name": p.name, "in_count": cnt})
    people_overlap.sort(key=lambda x: x["in_count"], reverse=True)
    return {
        "kind": "concept",
        "slug": slug,
        "name": c.name,
        "parents": list(c.parent_concepts),
        "children": list(c.child_concepts),
        "related": list(c.related_concepts),
        "contributing_sources": contributing[:limit],
        "central_people": people_overlap[:limit],
    }


def _dump_review(r, full: bool = False) -> dict:
    out = {
        "id": r.id,
        "kind": r.kind,
        "title": r.title,
        "status": r.status,
        "created_at": r.created_at.isoformat(),
    }
    if r.reviewed_at:
        out["reviewed_at"] = r.reviewed_at.isoformat()
    if full:
        out["preview"] = r.preview
        out["proposed_action"] = r.proposed_action
    return out


def _render_review_pending(config: Config, state) -> None:
    """Render /wiki/_review/pending.md from current review_queue state."""
    review_dir = config.vault_root / "wiki" / "_review"
    review_dir.mkdir(parents=True, exist_ok=True)
    path = review_dir / "pending.md"

    pending = [r for r in state.review_queue if r.status == "pending"]
    pending.sort(key=lambda r: r.created_at, reverse=True)

    lines = ["# Pending Reviews\n"]
    if not pending:
        lines.append(
            "_(nothing to review right now)_\n\n"
            "When the system has Claude-proposed actions waiting for your "
            "approval (concept page refreshes, hierarchy suggestions, Drive "
            "ingest candidates), they'll show up here.\n"
        )
    else:
        lines.append(
            f"_{len(pending)} item{'s' if len(pending) != 1 else ''} waiting "
            "for your decision. To approve / reject in chat, say:_\n"
            "- *approve <id>* / *reject <id>*\n"
            "- *approve all concept refreshes* / *reject all drive borderline*\n\n"
        )

        # Group by kind
        from collections import defaultdict
        by_kind: dict[str, list] = defaultdict(list)
        for r in pending:
            by_kind[r.kind].append(r)

        for kind, items in by_kind.items():
            lines.append(f"## {kind} ({len(items)})\n")
            for r in items:
                lines.append(f"### {r.title}")
                lines.append(f"_id: `{r.id}` · created {r.created_at.date().isoformat()}_\n")
                lines.append(r.preview)
                lines.append("")

    path.write_text("\n".join(lines))


# Dispatch table for executing review-queue proposed actions.
# Maps tool_name → callable that takes (config, args) and returns a result dict.
# Populated lazily inside _dispatch_review_action so we don't need to import
# every dependency at module load time.
def _dispatch_review_action(config: Config, tool_name: str, args: dict) -> dict:
    """Dispatch a proposed-action to the named tool. Used by approve_review."""
    if not tool_name:
        return {"error": "proposed_action missing 'tool' field"}

    # Build server fresh to access tools — heavy but only on approval.
    # Long-running approve is fine since user is reviewing in chat.
    if tool_name == "record_concept_page":
        from deep_reader.state import GlobalState
        state = GlobalState.load(config.state_file)
        return _do_record_concept_page(config, state, **args)
    if tool_name == "link_concepts":
        from deep_reader.state import GlobalState
        state = GlobalState.load(config.state_file)
        return _do_link_concepts(config, state, **args)
    if tool_name == "ingest_drive_doc":
        # The actual record_doc tool is on the MCP server. We could call
        # it via the server's call_tool API, but here we'll just record
        # the metadata and let Claude follow up with the real ingest.
        return {
            "note": (
                "ingest_drive_doc approved. Claude should now call "
                "record_doc with the structured analysis it prepared."
            ),
            "args": args,
        }
    return {"error": f"no dispatch handler for tool '{tool_name}'"}


def _do_record_concept_page(config: Config, state, **kwargs) -> dict:
    """Concrete handler for record_concept_page. Used by approve_review and
    callable directly from the tool of the same name."""
    from deep_reader.markdown import format_frontmatter
    from deep_reader.state import Concept

    name = kwargs.get("name", "").strip()
    if not name:
        return {"error": "name required"}
    slug = name.lower().replace(" ", "-")
    definition = (kwargs.get("definition") or "").strip()
    distillation = (kwargs.get("distillation") or "").strip()
    tensions = (kwargs.get("tensions") or "").strip()
    contributing = kwargs.get("contributing_sources") or []
    parent_concepts = kwargs.get("parent_concepts") or []
    child_concepts = kwargs.get("child_concepts") or []
    related_concepts = kwargs.get("related_concepts") or []

    # Update Concept state record
    if slug not in state.concepts:
        state.concepts[slug] = Concept(slug=slug, name=name)
    concept = state.concepts[slug]
    concept.parent_concepts = list(set(concept.parent_concepts + parent_concepts))
    concept.child_concepts = list(set(concept.child_concepts + child_concepts))
    concept.related_concepts = list(set(concept.related_concepts + related_concepts))
    concept.last_refreshed = datetime.now()
    concept.sources_at_last_refresh = len(contributing)
    state.save(config.state_file)

    # Render the page
    fm = {"name": name, "slug": slug, "type": "concept", "last_refreshed": datetime.now().date().isoformat()}
    parts = [format_frontmatter(fm), f"# {name}\n"]
    if definition:
        parts.append(f"## Definition\n{definition}\n")
    if parent_concepts or child_concepts or related_concepts:
        h_lines = ["## Hierarchy"]
        if parent_concepts:
            h_lines.append("**Parents:** " + ", ".join(f"[[concepts/{c}|{c}]]" for c in parent_concepts))
        if child_concepts:
            h_lines.append("**Children:** " + ", ".join(f"[[concepts/{c}|{c}]]" for c in child_concepts))
        if related_concepts:
            h_lines.append("**Related:** " + ", ".join(f"[[concepts/{c}|{c}]]" for c in related_concepts))
        parts.append("\n".join(h_lines) + "\n")
    if distillation:
        parts.append(f"## Distillation\n{distillation}\n")
    if tensions:
        parts.append(f"## Tensions / open questions\n{tensions}\n")
    if contributing:
        contrib_lines = ["## Contributing sources"]
        for slug_ in contributing:
            contrib_lines.append(f"- [[sources/{slug_}/_overview|{slug_}]]")
        parts.append("\n".join(contrib_lines) + "\n")

    path = config.wiki_concepts / f"{slug}.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text("\n".join(parts))
    return {"name": name, "slug": slug, "path": str(path.relative_to(config.vault_root))}


def _do_link_concepts(config: Config, state, parent: str, child: str, kind: str = "parent") -> dict:
    """Establish a hierarchy relationship between two concepts."""
    from deep_reader.state import Concept
    parent_slug = parent.strip().lower().replace(" ", "-")
    child_slug = child.strip().lower().replace(" ", "-")
    if parent_slug not in state.concepts:
        state.concepts[parent_slug] = Concept(slug=parent_slug, name=parent)
    if child_slug not in state.concepts:
        state.concepts[child_slug] = Concept(slug=child_slug, name=child)
    p = state.concepts[parent_slug]
    c = state.concepts[child_slug]
    if kind == "parent":
        if child_slug not in p.child_concepts:
            p.child_concepts.append(child_slug)
        if parent_slug not in c.parent_concepts:
            c.parent_concepts.append(parent_slug)
    elif kind == "related":
        if child_slug not in p.related_concepts:
            p.related_concepts.append(child_slug)
        if parent_slug not in c.related_concepts:
            c.related_concepts.append(parent_slug)
    state.save(config.state_file)
    return {"parent": parent_slug, "child": child_slug, "kind": kind, "linked": True}


def _rerender_after_action_change(config: Config, state, item) -> None:
    """Re-render every view that depends on a single action item.

    Called after add_action_item / add_waiting_on / close_action_item. Without
    this, per-person pages go stale relative to the central action_items.md
    and waiting_on.md files — e.g. closing a waiting-on item would update
    waiting_on.md but leave the owner's person page showing it as still open.
    """
    from deep_reader.wiki import Wiki, render_action_items, render_waiting_on
    from deep_reader.steps import people as people_step

    wiki = Wiki(config)
    render_action_items(wiki, state)
    render_waiting_on(wiki, state)

    # Re-render person pages that reference this item.
    slugs_to_refresh: set[str] = set()
    if item is not None:
        slugs_to_refresh.add(item.owner)
    # Vault owner page shows "My open action items" — refresh whenever any
    # mine item is touched (owner could be the vault owner slug, but the
    # match by name/email is authoritative).
    for p in state.people.values():
        if state.owner.matches(p.name) or state.owner.matches(p.email or ""):
            slugs_to_refresh.add(p.slug)

    for slug in slugs_to_refresh:
        if slug in state.people:
            people_step.render_person_page(state.people[slug], state, config.wiki_people)


def _do_structured_record(
    config: Config,
    source_type: str,
    title: str,
    body: str,
    date: str | None,
    attendees: list[str] | list[dict],
    result_payload: dict,
) -> dict:
    """Shared implementation for record_meeting / record_note / record_doc.

    No LLM. Writes raw file, builds a Source, seeds SourceState threads from
    the global pool, then calls reader.record_structured_source() to persist
    everything (overview, detail page, people, action items, threads).
    """
    from datetime import date as _date
    from deep_reader.markdown import format_frontmatter
    from deep_reader.reader import record_structured_source
    from deep_reader.sources.base import Source, SourceType
    from deep_reader.state import GlobalState, SourceState
    from deep_reader.wiki import Wiki

    type_enum = {
        "meeting": SourceType.MEETING,
        "note": SourceType.NOTE,
        "doc": SourceType.DOC,
    }[source_type]

    # Parse date string
    meeting_date = None
    if date:
        try:
            parts = date.split("-")
            meeting_date = _date(int(parts[0]), int(parts[1]), int(parts[2]))
        except Exception:
            pass

    # Extract attendee names for the Source
    attendee_names: list[str] = []
    for a in attendees or []:
        if isinstance(a, str):
            attendee_names.append(a)
        elif isinstance(a, dict) and a.get("name"):
            attendee_names.append(a["name"])

    # Write raw body to disk with frontmatter
    fm: dict = {"title": title, "type": source_type}
    if date:
        fm["date"] = date
    if attendee_names:
        fm["attendees"] = attendee_names

    slug = _slugify(title) or "untitled"
    raw_dir = _raw_dir_for(config, source_type)
    raw_dir.mkdir(parents=True, exist_ok=True)
    filename = f"{date}-{slug}.md" if date and source_type == "meeting" else f"{slug}.md"
    raw_path = raw_dir / filename
    # If the file already exists (re-record), rename with a suffix
    if raw_path.exists():
        import time
        filename = f"{raw_path.stem}-{int(time.time())}{raw_path.suffix}"
        raw_path = raw_dir / filename
    raw_path.write_text(format_frontmatter(fm) + body)

    # Build Source
    source = Source(
        path=raw_path,
        title=title,
        author=source_type,
        source_type=type_enum,
        meeting_date=meeting_date,
        attendees=attendee_names,
    )

    # Load state and set up SourceState, seeding threads from global
    state = GlobalState.load(config.state_file)
    wiki = Wiki(config)
    wiki.init_source(source.slug, source.title, source.author, source.source_type.value)

    if source.slug not in state.sources:
        state.sources[source.slug] = SourceState(
            source_slug=source.slug,
            source_path=str(source.path),
            threads=list(state.global_threads),
            source_type=source_type,
            meeting_date=date,
            attendees=attendee_names,
        )
    source_state = state.sources[source.slug]

    # Ensure result_payload has expected shape
    result_payload.setdefault("full_text", body)

    summary = record_structured_source(
        source, result_payload, source_state, state, wiki, config
    )
    summary["raw_file"] = str(raw_path.relative_to(config.vault_root))
    return summary




def _dump_action(a, state: GlobalState | None = None) -> dict:
    d = {
        "id": a.id, "description": a.description, "owner": a.owner,
        "source": a.source, "created_at": a.created_at.isoformat(),
        "status": a.status, "category": a.category,
    }
    if a.completed_at:
        d["completed_at"] = a.completed_at.isoformat()
    if state and a.owner in state.people:
        d["owner_name"] = state.people[a.owner].name
    return d


def _slugify(text: str) -> str:
    # Delegate to the hardened canonical slugify — handles em-dashes, curly
    # quotes, escape-sequence residue (e.g. literal "\u2014" sneaking in).
    from deep_reader.markdown import slugify
    return slugify(text)


def _auto_detect_type_path(path: Path) -> str:
    import re
    name = path.name.lower()
    if re.match(r"^\d{4}-\d{2}-\d{2}[-_\s]", name):
        return "meeting"
    if any(k in name for k in ["meeting", "1-1", "standup", "sync", "call"]):
        return "meeting"
    try:
        if path.suffix.lower() in {".md", ".txt"} and path.stat().st_size < 4000:
            return "note"
    except OSError:
        pass
    return "doc"


def _raw_dir_for(config: Config, source_type: str) -> Path:
    return {
        "book": config.raw_books,
        "article": config.raw_articles,
        "paper": config.raw_papers,
        "meeting": config.raw_meetings,
        "doc": config.raw_docs,
        "note": config.raw_notes,
    }.get(source_type, config.raw_docs)


def _ingest_path(config: Config, path: Path, source_type: str) -> None:
    """Extract text if needed; here we're mainly doing file conversion for non-md/txt."""
    if path.suffix.lower() in {".md", ".txt"}:
        return
    from deep_reader.sources.text import extract_text
    from deep_reader.sources.pdf import extract_pdf
    if path.suffix.lower() == ".pdf":
        text = extract_pdf(path)
    elif path.suffix.lower() == ".docx":
        try:
            import docx
        except ImportError:
            raise RuntimeError("python-docx not installed")
        doc = docx.Document(str(path))
        text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    else:
        text = extract_text(path)
    # Replace the file with a .md version for consistent downstream handling
    md_path = path.with_suffix(".md")
    md_path.write_text(text)
    if md_path != path:
        path.unlink()


def _read_new_source(config: Config, path: Path, source_type: str) -> None:
    """Invoke the read pipeline on a freshly-ingested source."""
    from deep_reader.sources.base import Source, SourceType
    from deep_reader.reader import read_source
    from deep_reader.markdown import parse_frontmatter
    from datetime import date as _date

    text = path.read_text()
    fm, _body = parse_frontmatter(text)
    title = fm.get("title") or path.stem
    meeting_date = None
    attendees: list[str] = []
    if source_type == "meeting":
        if fm.get("date"):
            try:
                parts = str(fm["date"]).split("-")
                meeting_date = _date(int(parts[0]), int(parts[1]), int(parts[2]))
            except Exception:
                pass
        if isinstance(fm.get("attendees"), list):
            attendees = fm["attendees"]

    stype_enum = {
        "book": SourceType.BOOK,
        "article": SourceType.ARTICLE,
        "paper": SourceType.PAPER,
        "meeting": SourceType.MEETING,
        "doc": SourceType.DOC,
        "note": SourceType.NOTE,
    }[source_type]

    source = Source(
        path=path,
        title=title,
        author={"meeting": "meeting", "doc": "doc", "note": "note"}.get(source_type, "Unknown"),
        source_type=stype_enum,
        meeting_date=meeting_date,
        attendees=attendees,
    )
    read_source(source, config, claude_code_llm, verbose=False, dry_run=False)


def main():
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("--vault", default="vault", help="Vault root directory")
    args = parser.parse_args()
    server = build_server(Path(args.vault))
    server.run()


if __name__ == "__main__":
    main()
